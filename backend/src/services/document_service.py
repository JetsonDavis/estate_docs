"""Service layer for document generation and merge operations."""

from sqlalchemy.orm import Session
from typing import Optional, Tuple, List
from fastapi import HTTPException, status
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import re
from html.parser import HTMLParser
from .s3_service import s3_service

_logger = logging.getLogger(__name__)


class HTMLToWordConverter(HTMLParser):
    """Custom HTML parser that converts HTML to Word document with proper formatting."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.style_stack = []  # Stack to track nested formatting
        self.list_level = 0
        self.in_list = False

    @staticmethod
    def _default_paragraph_spacing(paragraph):
        """
        Clear Word's default Normal-style space before/after and use single line spacing.

        Merged templates control vertical gaps with <cr> (line breaks), not empty paragraphs.
        Without this, centered titles and body text look like double-spaced paragraphs.
        """
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag == 'pagebreak':
            # Insert a Word page break
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
                self._default_paragraph_spacing(self.current_paragraph)
            run = self.current_paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            self.current_paragraph = None
            self.current_run = None

        elif tag == 'p':
            # Create new paragraph
            self.current_paragraph = self.doc.add_paragraph()
            self._default_paragraph_spacing(self.current_paragraph)
            self.current_run = None

            # Apply paragraph-level styling from style attribute
            if 'style' in attrs_dict:
                self._apply_paragraph_style(attrs_dict['style'])

            # Handle Quill class-based alignment and indentation
            if 'class' in attrs_dict:
                classes = attrs_dict['class'].split()
                if 'ql-align-center' in classes:
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'ql-align-right' in classes:
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'ql-align-justify' in classes:
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Handle Quill indent classes (ql-indent-1 through ql-indent-8)
                for cls in classes:
                    if cls.startswith('ql-indent-'):
                        try:
                            indent_level = int(cls.replace('ql-indent-', ''))
                            # Each indent level is approximately 0.5 inches
                            self.current_paragraph.paragraph_format.left_indent = Inches(indent_level * 0.5)
                        except ValueError:
                            pass

        elif tag == 'br':
            # Add line break within current paragraph
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
                self._default_paragraph_spacing(self.current_paragraph)
            if self.current_run is None:
                self.current_run = self.current_paragraph.add_run()
            self.current_run.add_break()

        elif tag in ('strong', 'b'):
            style_dict = {'bold': True}
            # Parse inline styles on strong tags too
            if 'style' in attrs_dict:
                style_dict.update(self._parse_inline_style(attrs_dict['style']))
            self.style_stack.append(style_dict)

        elif tag in ('em', 'i'):
            style_dict = {'italic': True}
            # Parse inline styles on em tags too
            if 'style' in attrs_dict:
                style_dict.update(self._parse_inline_style(attrs_dict['style']))
            self.style_stack.append(style_dict)

        elif tag == 'u':
            style_dict = {'underline': True}
            # Parse inline styles on u tags too
            if 'style' in attrs_dict:
                style_dict.update(self._parse_inline_style(attrs_dict['style']))
            self.style_stack.append(style_dict)

        elif tag == 'span':
            # Parse inline styles
            style_dict = {}
            if 'style' in attrs_dict:
                style_dict = self._parse_inline_style(attrs_dict['style'])

            # Handle Quill class-based font sizes (ql-size-*)
            if 'class' in attrs_dict:
                classes = attrs_dict['class'].split()
                for cls in classes:
                    if cls.startswith('ql-size-'):
                        size_str = cls.replace('ql-size-', '')
                        # Quill uses size values like '10px', '12px', etc.
                        match = re.match(r'(\d+)px', size_str)
                        if match:
                            style_dict['font_size'] = float(match.group(1))

            self.style_stack.append(style_dict)

        elif tag in ('ul', 'ol'):
            self.in_list = True
            self.list_level += 1

        elif tag == 'li':
            # Create list item paragraph
            self.current_paragraph = self.doc.add_paragraph(style='List Bullet' if self.in_list else None)
            self._default_paragraph_spacing(self.current_paragraph)
            self.current_run = None

    def handle_endtag(self, tag):
        if tag == 'p':
            self.current_paragraph = None
            self.current_run = None

        elif tag in ('strong', 'b', 'em', 'i', 'u', 'span'):
            if self.style_stack:
                self.style_stack.pop()

        elif tag in ('ul', 'ol'):
            self.list_level -= 1
            if self.list_level == 0:
                self.in_list = False

        elif tag == 'li':
            self.current_paragraph = None
            self.current_run = None

    def handle_data(self, data):
        # Skip completely empty data, but preserve whitespace-only if it contains tabs
        if not data.strip() and '\t' not in data:
            return

        # Ensure we have a paragraph
        if self.current_paragraph is None:
            self.current_paragraph = self.doc.add_paragraph()
            self._default_paragraph_spacing(self.current_paragraph)

        # Handle tabs by converting them to proper Word tab stops
        # Split data by tabs and add tab characters
        if '\t' in data:
            parts = data.split('\t')
            for i, part in enumerate(parts):
                if part:  # Add text part
                    self.current_run = self.current_paragraph.add_run(part)
                    self._apply_run_styles(self.current_run)
                if i < len(parts) - 1:  # Add tab between parts (not after last)
                    self.current_run = self.current_paragraph.add_run('\t')
                    self._apply_run_styles(self.current_run)
        else:
            # Create a new run with current styling
            self.current_run = self.current_paragraph.add_run(data)
            self._apply_run_styles(self.current_run)

    def _apply_run_styles(self, run):
        """Apply accumulated styles from stack to a run."""
        for style_dict in self.style_stack:
            if 'bold' in style_dict:
                run.bold = style_dict['bold']
            if 'italic' in style_dict:
                run.italic = style_dict['italic']
            if 'underline' in style_dict:
                run.underline = style_dict['underline']
            if 'font_size' in style_dict:
                font_size = style_dict['font_size']
                run.font.size = Pt(font_size)
                # Debug: write to file
                try:
                    with open('/tmp/quill_html_debug.html', 'a') as f:
                        f.write(f"\nApplying font size: {font_size}pt to text\n")
                except:
                    pass
            if 'color' in style_dict:
                run.font.color.rgb = style_dict['color']

    def _parse_inline_style(self, style_str):
        """Parse inline CSS style string and return dict of applicable styles."""
        style_dict = {}

        # Split by semicolon and parse each property
        for prop in style_str.split(';'):
            if ':' not in prop:
                continue
            key, value = prop.split(':', 1)
            key = key.strip().lower()
            value = value.strip()

            if key == 'font-size':
                # Extract numeric value from font-size (e.g., "14px" -> 14, "32px" -> 32)
                # Remove 'px', 'pt', or other units
                value_clean = value.replace('px', '').replace('pt', '').strip()
                try:
                    font_size_val = float(value_clean)
                    style_dict['font_size'] = font_size_val
                    # Debug: write to file
                    try:
                        with open('/tmp/quill_html_debug.html', 'a') as f:
                            f.write(f"\nParsed font-size: {value} -> {font_size_val}\n")
                    except:
                        pass
                except ValueError:
                    pass

            elif key == 'color':
                # Parse color (basic support for hex colors)
                if value.startswith('#'):
                    try:
                        rgb = RGBColor(
                            int(value[1:3], 16),
                            int(value[3:5], 16),
                            int(value[5:7], 16)
                        )
                        style_dict['color'] = rgb
                    except:
                        pass

        return style_dict

    def _apply_paragraph_style(self, style_str):
        """Apply paragraph-level styles like alignment, spacing, indentation."""
        if not self.current_paragraph:
            return

        for prop in style_str.split(';'):
            if ':' not in prop:
                continue
            key, value = prop.split(':', 1)
            key = key.strip().lower()
            value = value.strip()

            if key == 'text-align':
                if value == 'center':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif value == 'right':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif value == 'justify':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif key == 'margin-left' or key == 'padding-left':
                # Convert px to inches (rough approximation: 96px = 1 inch)
                match = re.match(r'(\d+(?:\.\d+)?)', value)
                if match:
                    px = float(match.group(1))
                    self.current_paragraph.paragraph_format.left_indent = Inches(px / 96)


from ..models.document import GeneratedDocument
from ..models.template import Template
from ..models.session import InputForm, SessionAnswer
from ..models.question import Question
from ..models.person import Person
from ..schemas.document import GenerateDocumentRequest


class DocumentService:
    """Service for document generation and merge operations."""

    @staticmethod
    def generate_document(
        db: Session,
        request: GenerateDocumentRequest,
        user_id: int
    ) -> GeneratedDocument:
        """
        Generate a document by merging session answers into a template.

        Args:
            db: Database session
            request: Document generation request
            user_id: User ID generating the document

        Returns:
            Generated document
        """
        # Get template
        template = db.query(Template).filter(
            Template.id == request.template_id,
            Template.is_active == True
        ).first()

        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )

        # Get session (verify user owns it)
        session = db.query(InputForm).filter(
            InputForm.id == request.session_id,
            InputForm.user_id == user_id
        ).first()

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found"
            )

        # Get all answers with their questions in a single joined query
        answer_pairs = db.query(SessionAnswer, Question).join(
            Question, SessionAnswer.question_id == Question.id
        ).filter(
            SessionAnswer.session_id == request.session_id
        ).all()

        # Build answer map: identifier -> formatted answer_value
        answer_map = DocumentService._build_answer_map(answer_pairs)

        # Build raw answer map (unformatted) so FOR EACH can parse JSON arrays
        raw_answer_map = DocumentService._build_raw_answer_map(answer_pairs)

        # Build conjunction info for repeatable groups
        conj_map, id_grp_map = DocumentService._build_conjunction_info(answer_pairs)

        # Merge template with answers
        merged_content = DocumentService._merge_template(
            template.markdown_content,
            answer_map,
            raw_answer_map,
            conj_map,
            id_grp_map
        )

        # Generate document name if not provided
        document_name = request.document_name or f"{template.name} - {session.client_identifier}"

        # Create generated document with placeholder S3 key
        document = GeneratedDocument(
            session_id=request.session_id,
            template_id=request.template_id,
            document_name=document_name,
            s3_key="",  # Will be updated after upload
            markdown_content=None,  # No longer storing in DB
            generated_by=user_id,
            generated_at=datetime.utcnow()
        )

        db.add(document)
        db.flush()  # Get the document ID without committing

        # Upload markdown to S3
        s3_key = s3_service.upload_markdown(merged_content, document.id, user_id)
        document.s3_key = s3_key

        db.commit()
        db.refresh(document)

        # Attach merged content to the response object so the API returns it
        # (it's stored in S3, not the DB, but callers expect it in the response)
        document.markdown_content = merged_content

        return document

    @staticmethod
    def _build_answer_map(answer_question_pairs: List[Tuple[SessionAnswer, Question]]) -> dict:
        """
        Build a map of question identifiers to answer values.

        Args:
            answer_question_pairs: List of (SessionAnswer, Question) tuples from a joined query

        Returns:
            Dictionary mapping identifiers to answer values
        """
        answer_map = {}

        for answer, question in answer_question_pairs:
            # Format person answers with conjunctions
            formatted_value = DocumentService._format_answer_value(
                answer.answer_value,
                question.question_type
            )
            # Store under lowercased full namespaced identifier (e.g., "group.poa_sign_date")
            answer_map[question.identifier.lower()] = formatted_value
            # Also store under stripped identifier (e.g., "poa_sign_date")
            # so templates can reference identifiers without namespace prefix
            if '.' in question.identifier:
                stripped = question.identifier.split('.', 1)[1].lower()
                # Only set stripped key if not already taken (first writer wins)
                if stripped not in answer_map:
                    answer_map[stripped] = formatted_value

        return answer_map

    @staticmethod
    def _build_raw_answer_map(answer_question_pairs: List[Tuple[SessionAnswer, Question]]) -> dict:
        """
        Build a map of question identifiers to raw (unformatted) answer values.

        Raw values preserve JSON arrays so FOR EACH loops can parse them.

        Args:
            answer_question_pairs: List of (SessionAnswer, Question) tuples from a joined query

        Returns:
            Dictionary mapping identifiers to raw answer values
        """
        raw_answer_map = {}

        for answer, question in answer_question_pairs:
            raw_answer_map[question.identifier.lower()] = answer.answer_value
            if '.' in question.identifier:
                stripped = question.identifier.split('.', 1)[1].lower()
                if stripped not in raw_answer_map:
                    raw_answer_map[stripped] = answer.answer_value

        return raw_answer_map

    @staticmethod
    def _format_answer_value(answer_value: str, question_type: str) -> str:
        """
        Format an answer value for display in merged documents.

        For person-type questions, converts JSON array with conjunctions to readable text.
        E.g., [{"name": "John", "conjunction": "and"}, {"name": "Jane"}] -> "John and Jane"

        Args:
            answer_value: Raw answer value from database
            question_type: Type of the question

        Returns:
            Formatted answer string
        """
        # Format date values as "Month Day, Year" (e.g., "March 9, 2026")
        if question_type == 'date':
            # Check if it's an array of dates (for repeatable questions)
            try:
                parsed = json.loads(answer_value)
                if isinstance(parsed, list):
                    # Format each date in the array
                    formatted_dates = []
                    for date_str in parsed:
                        try:
                            dt = datetime.strptime(date_str, '%Y-%m-%d')
                            formatted_dates.append(dt.strftime('%B %-d, %Y'))
                        except (ValueError, TypeError):
                            formatted_dates.append(date_str)
                    # Join with conjunctions for display (e.g., "March 10, 2026 and March 14, 2026")
                    if len(formatted_dates) == 0:
                        return ''
                    if len(formatted_dates) == 1:
                        return formatted_dates[0]
                    if len(formatted_dates) == 2:
                        return f'{formatted_dates[0]} and {formatted_dates[1]}'
                    return ', '.join(formatted_dates[:-1]) + ', and ' + formatted_dates[-1]
            except (json.JSONDecodeError, TypeError):
                pass

            # Single date value
            try:
                dt = datetime.strptime(answer_value, '%Y-%m-%d')
                return dt.strftime('%B %-d, %Y')
            except (ValueError, TypeError):
                return answer_value

        if question_type not in ('person', 'person_backup'):
            return answer_value

        # Try to parse as JSON array of person objects
        try:
            parsed = json.loads(answer_value)

            # Single person object (not in array) — wrap it so the
            # same normalisation / formatting logic handles it.
            if isinstance(parsed, dict):
                parsed = [parsed]

            if isinstance(parsed, list) and len(parsed) > 0:
                # Normalise each element: if it's a JSON string, decode it
                normalised = []
                for item in parsed:
                    if isinstance(item, dict):
                        normalised.append(item)
                    elif isinstance(item, str):
                        try:
                            obj = json.loads(item)
                            if isinstance(obj, dict):
                                normalised.append(obj)
                            elif isinstance(obj, list):
                                # Array-wrapped person, e.g. '[{"name":"James"}]'
                                for sub in obj:
                                    if isinstance(sub, dict):
                                        normalised.append(sub)
                                    else:
                                        normalised.append({'name': str(sub)})
                            else:
                                normalised.append({'name': item})
                        except (json.JSONDecodeError, TypeError):
                            normalised.append({'name': item})
                    else:
                        normalised.append({'name': str(item)})

                # Check if it's the format with objects containing name and conjunction
                if any('name' in p for p in normalised):
                    result_parts = []
                    for i, person in enumerate(normalised):
                        name = DocumentService._extract_plain_name(person.get('name', ''))
                        if not name:
                            continue
                        # The conjunction on THIS person indicates how it connects
                        # to the PREVIOUS person (e.g., "then Andrea" means
                        # Andrea follows the previous person with "then")
                        conjunction = person.get('conjunction', '') or 'and'
                        if i > 0:
                            if conjunction.lower() == 'then':
                                result_parts.append(', then')
                            else:
                                result_parts.append(conjunction)
                        result_parts.append(name)
                    return ' '.join(result_parts)

            return answer_value
        except (json.JSONDecodeError, TypeError):
            # Not JSON, return as-is
            return answer_value

    @staticmethod
    def _build_conjunction_info(answer_question_pairs: List[Tuple[SessionAnswer, Question]]) -> tuple:
        """
        Build conjunction info from person-type repeatable answers.

        Args:
            answer_question_pairs: List of (SessionAnswer, Question) tuples from a joined query

        Returns:
            Tuple of (conjunction_map, identifier_group_map) where:
            - conjunction_map: {repeatable_group_id: [conj1, conj2, ...]} from person entries
            - identifier_group_map: {identifier: repeatable_group_id} for all repeatable questions
        """
        conjunction_map = {}  # repeatable_group_id -> [conjunctions]
        identifier_group_map = {}  # identifier -> repeatable_group_id

        for answer, question in answer_question_pairs:
            # Map every repeatable question's identifier to its group
            if question.repeatable and question.repeatable_group_id:
                ident = question.identifier.lower()
                identifier_group_map[ident] = question.repeatable_group_id
                if '.' in question.identifier:
                    stripped = question.identifier.split('.', 1)[1].lower()
                    identifier_group_map[stripped] = question.repeatable_group_id

                # Extract conjunctions from person-type answers
                if question.question_type in ('person', 'person_backup'):
                    group_id = question.repeatable_group_id
                    if group_id not in conjunction_map:
                        try:
                            parsed = json.loads(answer.answer_value)
                            if isinstance(parsed, list):
                                conjunctions = []
                                for item in parsed:
                                    decoded = item
                                    if isinstance(decoded, str):
                                        try:
                                            decoded = json.loads(decoded)
                                        except (json.JSONDecodeError, TypeError):
                                            pass
                                    # Unwrap array-wrapped person, e.g. [{"name":"James"}]
                                    if isinstance(decoded, list) and len(decoded) == 1 and isinstance(decoded[0], dict):
                                        decoded = decoded[0]
                                    if isinstance(decoded, dict):
                                        conjunctions.append(decoded.get('conjunction', 'and'))
                                    else:
                                        conjunctions.append('and')
                                conjunction_map[group_id] = conjunctions
                        except (json.JSONDecodeError, TypeError):
                            pass

        return conjunction_map, identifier_group_map

    @staticmethod
    def _is_value_empty(value: str) -> bool:
        """Check if a value should be considered empty."""
        if not value:
            return True
        if not value.strip():
            return True
        # Check for "[identifier: NOT ANSWERED]" pattern but not JSON arrays
        if value.startswith('[') and value.endswith(']') and ': NOT ANSWERED]' in value:
            return True
        return False

    # ── Shared word lists for counter tokens ─────────────────────────
    _CARDINAL_WORDS = [
        '', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven',
        'Eight', 'Nine', 'Ten', 'Eleven', 'Twelve', 'Thirteen',
        'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen',
        'Nineteen', 'Twenty'
    ]
    _ORDINAL_WORDS = [
        '', 'First', 'Second', 'Third', 'Fourth', 'Fifth', 'Sixth',
        'Seventh', 'Eighth', 'Ninth', 'Tenth', 'Eleventh', 'Twelfth',
        'Thirteenth', 'Fourteenth', 'Fifteenth', 'Sixteenth',
        'Seventeenth', 'Eighteenth', 'Nineteenth', 'Twentieth'
    ]

    # ── Pre-compiled regexes ──────────────────────────────────────────
    _FOREACH_OPEN_RE = re.compile(
        r'\{\{\s*(?:FOR\s+EACH|FOREACH)(?:\((\d+)\))?\s+(?:<<)?([^>=!\s\}>]+)(?:>>)?'
        r'(?:\s+WHERE\s+(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*(=|!=)\s*["\']([^"\']*)["\'])?'
        r'\s*\}\}',
        re.IGNORECASE
    )
    _FOREACH_CLOSE_RE = re.compile(
        r'\{\{\s*END\s+(?:FOR\s+EACH|FOREACH)\s*\}\}',
        re.IGNORECASE
    )
    _PEOPLELOOP_OPEN_RE = re.compile(
        r'\{\{\s*PEOPLELOOP\s+(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*\}\}',
        re.IGNORECASE
    )
    _PEOPLELOOP_CLOSE_RE = re.compile(
        r'\{\{\s*END\s+PEOPLELOOP\s*\}\}',
        re.IGNORECASE
    )
    # Matches any loop-opening tag (FOREACH or PEOPLELOOP) for placeholder protection
    _NESTED_LOOP_RE = re.compile(
        r'\{\{\s*(?:(?:FOR\s+EACH|FOREACH)(?:\(\d+\))?\s+|PEOPLELOOP\s+)',
        re.IGNORECASE
    )
    # Block-level regexes for depth counting across IF, FOREACH, and PEOPLELOOP
    _BLOCK_OPEN_RE = re.compile(
        r'\{\{\s*(?:(?:FOR\s+EACH|FOREACH)(?:\(\d+\))?\s+|PEOPLELOOP\s+|IF\s|SWITCH\s)',
        re.IGNORECASE
    )
    _BLOCK_CLOSE_RE = re.compile(
        r'\{\{\s*END(?:\s+(?:FOR\s+EACH|FOREACH|PEOPLELOOP|SWITCH))?\s*\}\}',
        re.IGNORECASE
    )
    _IF_OPEN_RE = re.compile(r'\{\{\s*IF\s+(.*?)\s*\}\}', re.IGNORECASE)
    _END_RE = re.compile(r'\{\{\s*END\s*\}\}', re.IGNORECASE)
    _ELSE_RE = re.compile(r'\{\{\s*ELSE\s*\}\}', re.IGNORECASE)
    _SWITCH_OPEN_RE = re.compile(
        r'\{\{\s*SWITCH\s+(?:[<«‹]{2})?([^>»›\}]+?)(?:[>»›]{2})?\s*\}\}',
        re.IGNORECASE
    )
    _CASE_RE = re.compile(
        r"""\{\{\s*CASE\s+(?:["'\u201c\u201d\u2018\u2019\u00ab\u00bb]([^"'\u201c\u201d\u2018\u2019\u00ab\u00bb]*?)["'\u201c\u201d\u2018\u2019\u00ab\u00bb]|(\S+?))\s*\}\}""",
        re.IGNORECASE
    )
    _END_SWITCH_RE = re.compile(r'\{\{\s*END\s+SWITCH\s*\}\}', re.IGNORECASE)
    _COUNTER_RE = re.compile(r'(###|##%|##A|##V|##)(?:\+(\d*))?')
    _COUNTER_RESET_RE = re.compile(r'#/A')
    _IDENTIFIER_RE = re.compile(r'[<«‹]{2}([^>»›]+?)(?:\[(\d+)\])?(?:\.([^>»›]+))?[>»›]{2}')
    _CONDITIONAL_RE = re.compile(r'\[\[(.*?)\]\]', re.DOTALL)

    @staticmethod
    def _parse_array(raw: str):
        """Try to parse a string as a JSON array; return list or None.

        Also wraps a single dict in a list so that person-type answers
        stored as ``{"name":"..."}`` (instead of ``[{"name":"..."}]``)
        are treated as one-element arrays.
        """
        if not raw:
            return None
        try:
            parsed = json.loads(raw)
            if isinstance(parsed, list):
                return parsed
            if isinstance(parsed, dict):
                return [parsed]
        except (json.JSONDecodeError, TypeError):
            pass
        return None

    @staticmethod
    def _extract_plain_name(name_value) -> str:
        """Extract a plain name string from a possibly JSON-encoded person object.

        Handles cases where a person's 'name' field is itself a JSON string
        like '{"name":"john sample"}' instead of just 'john sample'.
        """
        if not isinstance(name_value, str):
            return str(name_value) if name_value is not None else ''
        # Try decoding up to 3 levels of nesting
        current = name_value
        for _ in range(3):
            try:
                parsed = json.loads(current)
                if isinstance(parsed, dict):
                    inner = parsed.get('name')
                    if inner is not None:
                        current = str(inner)
                        continue
                    return str(parsed)
                break
            except (json.JSONDecodeError, TypeError):
                break
        return current

    @staticmethod
    def _decode_json_item(item):
        """Decode a possibly multi-level JSON-encoded item to its innermost dict or str."""
        decoded = item
        if isinstance(decoded, str):
            for _ in range(3):
                try:
                    obj = json.loads(decoded)
                    if isinstance(obj, dict):
                        return obj
                    elif isinstance(obj, list):
                        # Array-wrapped person data, e.g. '[{"name":"James"}]'
                        # Unwrap single-element lists to the inner dict/value
                        if len(obj) == 1 and isinstance(obj[0], dict):
                            return obj[0]
                        return obj
                    elif isinstance(obj, str):
                        decoded = obj
                    else:
                        break
                except (json.JSONDecodeError, TypeError):
                    break
        return decoded

    @staticmethod
    def _strip_value_html(value: str) -> str:
        """Strip HTML tags from an answer value to prevent unwanted formatting.

        Answer values stored by the Quill editor may contain <p>, <br>, and
        inline formatting tags.  These must be removed before substitution
        so they don't create spurious carriage returns in the output.
        """
        if not value:
            return value
        cleaned = re.sub(r'</?(?:p|br|span|strong|em|b|i|u)\b[^>]*/?>',
                         '', value)
        return cleaned.strip()

    @staticmethod
    def _format_item(item, field: str = None) -> str:
        """Format a single array element for output.

        Handles double-encoded JSON strings where each element may be
        a JSON string containing a JSON object (e.g. '{"name":"Alice"}').
        """
        decoded = DocumentService._decode_json_item(item)

        if isinstance(decoded, dict):
            if field:
                val = decoded.get(field, '')
                return DocumentService._extract_plain_name(str(val)) if val else ''
            name = decoded.get('name', str(decoded))
            return DocumentService._extract_plain_name(name)
        if isinstance(decoded, list):
            # Safety net for list values (e.g. multi-element array-wrapped person data)
            names = []
            for elem in decoded:
                if isinstance(elem, dict):
                    names.append(DocumentService._extract_plain_name(elem.get('name', '')))
                else:
                    names.append(str(elem))
            return ', '.join(n for n in names if n)
        if isinstance(decoded, str):
            if field:
                return DocumentService._strip_value_html(decoded) if field == 'name' else ''
            return DocumentService._strip_value_html(decoded)
        return DocumentService._strip_value_html(str(decoded))

    @staticmethod
    def _int_to_roman(n: int) -> str:
        """Convert a positive integer to an uppercase Roman numeral string."""
        if n <= 0:
            return 'I'
        vals = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I'),
        ]
        result = ''
        for value, numeral in vals:
            while n >= value:
                result += numeral
                n -= value
        return result

    @staticmethod
    def _counter_to_str(token: str, n: int) -> str:
        """Convert a counter token (##, ###, ##%, ##A, ##V) and number to its string representation."""
        if token == '###':
            return DocumentService._CARDINAL_WORDS[n] if n < len(DocumentService._CARDINAL_WORDS) else str(n)
        elif token == '##%':
            return DocumentService._ORDINAL_WORDS[n] if n < len(DocumentService._ORDINAL_WORDS) else f'{n}th'
        elif token == '##V':
            return DocumentService._int_to_roman(n)
        elif token == '##A':
            # Convert number to uppercase letter (1=A, 2=B, ..., 26=Z, 27=AA, 28=AB, etc.)
            if n <= 0:
                return 'A'
            result = ''
            n_temp = n
            while n_temp > 0:
                n_temp -= 1
                result = chr(65 + (n_temp % 26)) + result
                n_temp //= 26
            return result
        else:
            return str(n)

    @staticmethod
    def _process_foreach_blocks(text: str, answer_map: dict, raw_map: dict, global_counter: list, identifier_group_map: dict = None) -> str:
        """Process {{ FOR EACH identifier }} ... {{ END FOR EACH }} blocks.

        Uses depth-counting to find the matching closing tag so that
        nested FOREACH blocks are handled correctly.  After expanding the
        outer body, any inner FOREACH blocks are processed recursively.

        Args:
            text: Template text potentially containing FOR EACH blocks
            answer_map: Formatted identifier -> value map
            raw_map: Raw (unformatted) identifier -> value map
            global_counter: Mutable [int] shared counter for ## tokens
            identifier_group_map: {identifier: repeatable_group_id} so we only
                index into arrays for identifiers in the same repeatable group
                as the loop identifier; other-group identifiers use the
                conjunction-joined value from answer_map instead.

        Returns:
            Text with FOR EACH blocks expanded
        """
        id_grp = identifier_group_map or {}
        # Splits text into {{ }} tags, \x00 placeholders, and plain text
        _seg_split = re.compile(r'(\{\{.*?\}\}|\x00NF\d+\x00)')

        result = []
        pos = 0

        while pos < len(text):
            open_match = DocumentService._FOREACH_OPEN_RE.search(text, pos)
            if not open_match:
                result.append(text[pos:])
                break

            # Text before the opening tag
            result.append(text[pos:open_match.start()])

            counter_start_str = open_match.group(1)
            loop_identifier = open_match.group(2).lower()
            where_identifier = open_match.group(3)
            where_operator = open_match.group(4)
            where_value = open_match.group(5)
            body_start = open_match.end()

            # ── Depth-count to find matching close ──────────────────────
            # Track ALL block opens (IF + FOREACH) and closes (END, END
            # FOR EACH, END FOREACH) so that {{END}} used as a FOREACH
            # closer is handled correctly.
            depth = 1
            scan = body_start
            body_end = None
            close_end = None

            while depth > 0 and scan < len(text):
                next_open = DocumentService._BLOCK_OPEN_RE.search(text, scan)
                next_close = DocumentService._BLOCK_CLOSE_RE.search(text, scan)

                if next_close is None:
                    break

                if next_open and next_open.start() < next_close.start():
                    depth += 1
                    scan = next_open.end()
                else:
                    depth -= 1
                    if depth == 0:
                        body_end = next_close.start()
                        close_end = next_close.end()
                    else:
                        scan = next_close.end()

            if body_end is None:
                # No explicit matching close found.  Treat end-of-text as
                # an implicit close so the FOREACH still expands (handles
                # templates whose outer FOREACH omits its closing tag).
                _logger.debug(f"FOR EACH: no matching close for '{loop_identifier}', using end-of-text as implicit close")
                body_end = len(text)
                close_end = len(text)

            body_template = text[body_start:body_end]

            # ── Resolve loop array ─────────────────────────────────────
            if counter_start_str:
                global_counter[0] = int(counter_start_str) - 1

            raw_value = raw_map.get(loop_identifier, '') or answer_map.get(loop_identifier, '')
            loop_array = DocumentService._parse_array(raw_value)

            if not loop_array or len(loop_array) == 0:
                _logger.debug(f"FOR EACH: identifier '{loop_identifier}' has no array data, removing block")
                pos = close_end
                continue

            instance_count = len(loop_array)

            # ── Build WHERE filter ─────────────────────────────────────
            if where_identifier and where_operator and where_value is not None:
                filter_ident = where_identifier.strip('"\'').lower()
                filter_raw = raw_map.get(filter_ident, '') or answer_map.get(filter_ident, '')
                filter_array = DocumentService._parse_array(filter_raw)
                if filter_array is None:
                    filter_array = [filter_raw] * instance_count

                included_indices = []
                for i in range(instance_count):
                    item_val = ''
                    if i < len(filter_array):
                        item = filter_array[i]
                        if isinstance(item, str):
                            try:
                                decoded = json.loads(item)
                                if isinstance(decoded, dict):
                                    item_val = str(decoded.get('name', decoded.get('value', '')))
                                else:
                                    item_val = str(decoded)
                            except (json.JSONDecodeError, TypeError):
                                item_val = str(item)
                        elif isinstance(item, dict):
                            item_val = str(item.get('name', item.get('value', '')))
                        else:
                            item_val = str(item)

                    matches = item_val.strip().lower() == where_value.strip().lower()
                    if where_operator == '=' and matches:
                        included_indices.append(i)
                    elif where_operator == '!=' and not matches:
                        included_indices.append(i)

                _logger.debug(f"FOR EACH WHERE: '{filter_ident}' {where_operator} '{where_value}' -> indices {included_indices} of {instance_count}")
            else:
                included_indices = list(range(instance_count))

            _logger.debug(f"FOR EACH: iterating '{loop_identifier}' with {len(included_indices)} instances (of {instance_count} total)")

            body_identifiers_raw = re.findall(r'<<([^>]+)>>', body_template)

            # Filter out identifiers with explicit [N] array indexing (e.g. <<ident[1]>>).
            # These request a specific array element and must be left for Pass 5
            # (_replace_identifiers) which correctly parses the [N] index.
            _array_idx_re = re.compile(r'\[\d+\]')
            body_identifiers_raw = [i for i in body_identifiers_raw if not _array_idx_re.search(i)]

            # Determine the repeatable group of the loop identifier
            loop_group = id_grp.get(loop_identifier)

            identifier_arrays = {}
            same_group = set()
            for ident in body_identifiers_raw:
                base_ident = ident.split('.', 1)[0].lower() if '.' in ident else ident.lower()
                if base_ident not in identifier_arrays:
                    ident_group = id_grp.get(base_ident)
                    if loop_group and ident_group == loop_group:
                        raw = raw_map.get(base_ident, '') or answer_map.get(base_ident, '')
                        identifier_arrays[base_ident] = DocumentService._parse_array(raw)
                        same_group.add(base_ident)
                    elif ident_group and ident_group != loop_group:
                        identifier_arrays[base_ident] = None
                    else:
                        raw = raw_map.get(base_ident, '') or answer_map.get(base_ident, '')
                        arr = DocumentService._parse_array(raw)
                        identifier_arrays[base_ident] = arr
                        # Non-repeatable followups inside a repeatable parent store
                        # multi-instance data as arrays (combined by the frontend)
                        # but have no repeatable_group_id.  If the array length
                        # matches the loop array, treat as same-group so values
                        # are indexed per-iteration instead of shown as raw JSON.
                        if arr is not None and len(arr) == instance_count:
                            same_group.add(base_ident)

            # ── Expand body for each included index ────────────────────
            output_parts = []
            for idx in included_indices:
                instance_body = body_template

                # ── Protect nested loop blocks with placeholders ─────
                # MUST happen before counter replacement so that ##%
                # tokens inside nested FOREACH/PEOPLELOOP blocks are
                # preserved for their own loop processing.
                # Matches both FOREACH and PEOPLELOOP opens; uses
                # _BLOCK_OPEN/CLOSE_RE for depth counting so {{END}}
                # closers and nested IF blocks are tracked correctly.
                placeholders = {}
                protected = instance_body
                ph_idx = 0
                s_pos = 0
                while True:
                    i_open = DocumentService._NESTED_LOOP_RE.search(protected, s_pos)
                    if not i_open:
                        break
                    d = 1
                    s = i_open.end()
                    i_close_end = None
                    while d > 0 and s < len(protected):
                        n_o = DocumentService._BLOCK_OPEN_RE.search(protected, s)
                        n_c = DocumentService._BLOCK_CLOSE_RE.search(protected, s)
                        if n_c is None:
                            break
                        if n_o and n_o.start() < n_c.start():
                            d += 1
                            s = n_o.end()
                        else:
                            d -= 1
                            if d == 0:
                                i_close_end = n_c.end()
                            else:
                                s = n_c.end()
                    if i_close_end is not None:
                        ph = f'\x00NF{ph_idx}\x00'
                        placeholders[ph] = protected[i_open.start():i_close_end]
                        protected = protected[:i_open.start()] + ph + protected[i_close_end:]
                        ph_idx += 1
                        s_pos = i_open.start() + len(ph)
                    else:
                        s_pos = i_open.end()

                instance_body = protected

                # Only manage counter tokens when FOREACH explicitly
                # specifies a counter start (e.g. {{ FOR EACH 1 <<id>> }}).
                # Without a counter start, leave tokens for PEOPLELOOP / Pass 4.
                if counter_start_str:
                    global_counter[0] += 1

                    def _foreach_counter_replace(m, _gc=global_counter):
                        token = m.group(1)
                        plus_str = m.group(2)
                        inc = int(plus_str) if plus_str else 0
                        return DocumentService._counter_to_str(token, _gc[0] + inc)

                    instance_body = DocumentService._COUNTER_RE.sub(_foreach_counter_replace, instance_body)

                # Split into {{ }} tags, placeholders, and plain text
                segments = _seg_split.split(instance_body)

                for orig_ident in body_identifiers_raw:
                    ident = orig_ident.lower()
                    if '.' in ident:
                        base_ident, field_name = ident.split('.', 1)
                    else:
                        base_ident, field_name = ident, None

                    arr = identifier_arrays.get(base_ident)
                    if arr is not None and base_ident in same_group and idx < len(arr):
                        replacement = DocumentService._format_item(arr[idx], field_name)
                    elif arr is not None and base_ident in same_group:
                        replacement = ''
                    else:
                        scalar = answer_map.get(ident, '') or answer_map.get(base_ident, '')
                        replacement = DocumentService._strip_value_html(scalar) if not DocumentService._is_value_empty(scalar) else ''

                    target = f'<<{orig_ident}>>'
                    for j in range(len(segments)):
                        if not segments[j].startswith('{{') and not segments[j].startswith('\x00'):
                            segments[j] = segments[j].replace(target, replacement)

                instance_body = ''.join(segments)

                # ── Process IF blocks with per-iteration answer maps ──
                # MUST happen BEFORE restoring nested loop placeholders so
                # that IF blocks inside PEOPLELOOP/FOREACH bodies are
                # preserved for their own loop processing (e.g. IF TYPE()).
                # Same-group identifiers must resolve to their per-iteration
                # scalar values (not the full array) so that conditions like
                # {{IF <<amendment_type>> = "Update SSTEE"}} work correctly.
                #
                # iter_answer gets per-iteration scalars so that simple
                # identifier lookups (e.g. {{IF principal = "Jeff"}}) resolve
                # to the current person.  iter_raw is NOT overridden so that
                # subscripted lookups like principal[1] still access the
                # original full array.
                iter_answer = dict(answer_map)
                iter_raw = dict(raw_map)
                for orig_ident in body_identifiers_raw:
                    ident = orig_ident.lower()
                    base_ident = ident.split('.', 1)[0] if '.' in ident else ident
                    arr = identifier_arrays.get(base_ident)
                    if arr is not None and base_ident in same_group and idx < len(arr):
                        field_name = ident.split('.', 1)[1] if '.' in ident else None
                        val = DocumentService._format_item(arr[idx], field_name)
                        iter_answer[ident] = val
                # Always add the loop identifier itself to iter_answer with
                # the per-iteration scalar.  This is needed when the loop
                # identifier only appears in IF conditions (without <<>>
                # brackets) or only in subscripted form (<<principal[1]>>)
                # and therefore isn't in body_identifiers_raw.
                if idx < len(loop_array):
                    iter_answer[loop_identifier] = DocumentService._format_item(loop_array[idx])
                instance_body = DocumentService._process_if_blocks(
                    instance_body, iter_answer, iter_raw
                )
                instance_body = DocumentService._process_switch_blocks(
                    instance_body, iter_answer, iter_raw
                )

                # Restore nested FOREACH/PEOPLELOOP blocks AFTER IF processing
                for ph, original in placeholders.items():
                    instance_body = instance_body.replace(ph, original)

                output_parts.append(instance_body)

            expanded = ''.join(output_parts)

            # Recursively process nested FOREACH blocks
            expanded = DocumentService._process_foreach_blocks(
                expanded, answer_map, raw_map, global_counter, identifier_group_map
            )

            result.append(expanded)
            pos = close_end

        return ''.join(result)

    @staticmethod
    def _group_people_by_then(raw_json: str) -> list:
        """Split a person-type JSON array into groups separated by 'then' conjunctions.

        Each group is a list of dicts with 'name' and 'conjunction' keys.
        A new group starts whenever an entry (other than the first) has
        conjunction == 'then'.

        Returns:
            List of groups, where each group is a list of person dicts.
        """
        try:
            parsed = json.loads(raw_json)
        except (json.JSONDecodeError, TypeError):
            return []
        if not isinstance(parsed, list) or len(parsed) == 0:
            return []

        # Normalise each element
        normalised = []
        for item in parsed:
            decoded = DocumentService._decode_json_item(item)
            if isinstance(decoded, dict):
                normalised.append(decoded)
            elif isinstance(decoded, str):
                normalised.append({'name': decoded})
            else:
                normalised.append({'name': str(decoded)})

        groups: list = []
        current_group: list = []

        for i, person in enumerate(normalised):
            conj = (person.get('conjunction', '') or '').lower()
            if i > 0 and conj == 'then':
                # Close the current group and start a new one
                if current_group:
                    groups.append(current_group)
                current_group = []
            current_group.append(person)

        if current_group:
            groups.append(current_group)

        return groups

    @staticmethod
    def _format_people_group(group: list) -> str:
        """Join a group of person dicts into a single display string.

        Uses each person's conjunction ('and', 'or') to join names.
        The first person in the group has no preceding conjunction.
        """
        parts = []
        for j, person in enumerate(group):
            name = DocumentService._extract_plain_name(person.get('name', ''))
            if not name:
                continue
            if j > 0:
                conj = (person.get('conjunction', '') or 'and').lower()
                # 'then' shouldn't appear inside a group, but fall back to 'and'
                if conj == 'then':
                    conj = 'and'
                parts.append(conj)
            parts.append(name)
        return ' '.join(parts)

    @staticmethod
    def _format_people_group_field(group: list, arr: list, start_idx: int, field: str = None) -> str:
        """Join a same-group identifier's values using the group's conjunction pattern.

        Mirrors _format_people_group but extracts field values from `arr`
        instead of names from the group dicts.  For example, if group is
        [Mary(and), Sally] and arr has ['daughter', 'wife'], this returns
        'daughter and wife'.
        """
        parts = []
        for j, person in enumerate(group):
            idx = start_idx + j
            if idx < len(arr):
                val = DocumentService._format_item(arr[idx], field)
            else:
                val = ''
            if not val:
                continue
            if j > 0 and parts:
                conj = (person.get('conjunction', '') or 'and').lower()
                if conj == 'then':
                    conj = 'and'
                parts.append(conj)
            parts.append(val)
        return ' '.join(parts)

    @staticmethod
    def _get_group_type(group: list) -> str:
        """Determine the type of a PEOPLELOOP group.

        Returns:
            'single'    — one person in the group
            'joint-and' — multiple people joined by 'and'
            'joint-or'  — multiple people joined by 'or'
        """
        if len(group) <= 1:
            return 'single'
        # Check conjunctions within the group (persons after the first)
        for person in group[1:]:
            conj = (person.get('conjunction', '') or 'and').lower()
            if conj == 'or':
                return 'joint-or'
        return 'joint-and'

    @staticmethod
    def _process_peopleloop_blocks(text: str, answer_map: dict, raw_map: dict,
                                    global_counter: list,
                                    identifier_group_map: dict = None) -> str:
        """Process {{ PEOPLELOOP <<identifier>> }} ... {{ END PEOPLELOOP }} blocks.

        Groups a person-type repeatable field by 'then' conjunctions.
        Each group becomes one loop iteration where <<identifier>> is
        replaced with the names joined by 'and'/'or'.

        Same-group identifiers are resolved for the FIRST person index
        in each group so that fields like <<sstee_relation>> work.
        """
        id_grp = identifier_group_map or {}
        _seg_split = re.compile(r'(\{\{.*?\}\}|\x00NF\d+\x00)')
        _counter_reset_re = re.compile(r'\{\{\s*COUNTER\s+RESET\s*\}\}|##/', re.IGNORECASE)

        result = []
        pos = 0

        while pos < len(text):
            open_match = DocumentService._PEOPLELOOP_OPEN_RE.search(text, pos)
            if not open_match:
                # Process counter tokens and COUNTER RESET tags in remaining text
                remaining = text[pos:]
                if _counter_reset_re.search(remaining):
                    global_counter[0] = 0
                    remaining = _counter_reset_re.sub('', remaining)
                remaining = DocumentService._replace_counter_tokens(remaining, global_counter)
                result.append(remaining)
                break

            # Text before this PEOPLELOOP — process counter tokens and resets
            before_text = text[pos:open_match.start()]
            if _counter_reset_re.search(before_text):
                global_counter[0] = 0
                before_text = _counter_reset_re.sub('', before_text)
            before_text = DocumentService._replace_counter_tokens(before_text, global_counter)
            result.append(before_text)

            loop_identifier = open_match.group(1).lower()
            body_start = open_match.end()

            # ── Depth-count to find matching close ──────────────────────
            depth = 1
            scan = body_start
            body_end = None
            close_end = None

            while depth > 0 and scan < len(text):
                next_open = DocumentService._BLOCK_OPEN_RE.search(text, scan)
                next_close = DocumentService._BLOCK_CLOSE_RE.search(text, scan)

                if next_close is None:
                    break

                if next_open and next_open.start() < next_close.start():
                    depth += 1
                    scan = next_open.end()
                else:
                    depth -= 1
                    if depth == 0:
                        body_end = next_close.start()
                        close_end = next_close.end()
                    else:
                        scan = next_close.end()

            if body_end is None:
                body_end = len(text)
                close_end = len(text)

            body_template = text[body_start:body_end]

            # ── Get person array and group by 'then' ────────────────────
            raw_value = raw_map.get(loop_identifier, '') or ''
            groups = DocumentService._group_people_by_then(raw_value)

            if not groups:
                pos = close_end
                continue

            # Determine loop group for same-group resolution
            loop_group = id_grp.get(loop_identifier)

            # Collect all identifiers referenced in the body
            body_identifiers_raw = DocumentService._IDENTIFIER_RE.findall(body_template)
            # body_identifiers_raw is list of tuples: (ident, subscript, field)
            # Flatten to just the identifier names
            body_ident_names = set()
            for ident_tuple in body_identifiers_raw:
                base = ident_tuple[0].lower() if isinstance(ident_tuple, tuple) else ident_tuple.lower()
                body_ident_names.add(base)

            # Parse same-group arrays for parallel resolution
            identifier_arrays = {}
            same_group = set()
            for ident in body_ident_names:
                base_ident = ident.split('.', 1)[0] if '.' in ident else ident
                if base_ident not in identifier_arrays:
                    ident_group = id_grp.get(base_ident)
                    if loop_group and ident_group == loop_group:
                        raw = raw_map.get(base_ident, '') or answer_map.get(base_ident, '')
                        identifier_arrays[base_ident] = DocumentService._parse_array(raw)
                        same_group.add(base_ident)

            # ── Expand body for each people group ───────────────────────
            # Track the starting raw index for each group so parallel
            # same-group identifiers can be resolved.
            group_start_indices = []
            raw_idx = 0
            for group in groups:
                group_start_indices.append(raw_idx)
                raw_idx += len(group)

            output_parts = []
            for g_idx, group in enumerate(groups):
                instance_body = body_template

                # The joined name string for this group
                group_name = DocumentService._format_people_group(group)
                first_idx = group_start_indices[g_idx]

                # Split body into segments and replace identifiers
                segments = _seg_split.split(instance_body)

                for ident_tuple in body_identifiers_raw:
                    if isinstance(ident_tuple, tuple):
                        orig_ident = ident_tuple[0]
                        field_name = ident_tuple[2] if len(ident_tuple) > 2 and ident_tuple[2] else None
                    else:
                        orig_ident = ident_tuple
                        field_name = None

                    ident = orig_ident.lower()
                    base_ident = ident.split('.', 1)[0] if '.' in ident else ident

                    if base_ident == loop_identifier:
                        replacement = group_name
                    elif base_ident in same_group:
                        arr = identifier_arrays.get(base_ident)
                        if arr is not None and first_idx < len(arr):
                            fld = ident.split('.', 1)[1] if '.' in ident else field_name
                            replacement = DocumentService._format_people_group_field(
                                group, arr, first_idx, fld
                            )
                        else:
                            replacement = ''
                    else:
                        scalar = answer_map.get(ident, '') or answer_map.get(base_ident, '')
                        replacement = DocumentService._strip_value_html(scalar) if not DocumentService._is_value_empty(scalar) else ''

                    target = f'<<{orig_ident}>>'
                    for j in range(len(segments)):
                        if not segments[j].startswith('{{') and not segments[j].startswith('\x00'):
                            segments[j] = segments[j].replace(target, replacement)

                instance_body = ''.join(segments)

                # Process IF blocks with per-group answer maps
                iter_answer = dict(answer_map)
                iter_raw = dict(raw_map)
                iter_answer[loop_identifier] = group_name
                iter_raw[loop_identifier] = group_name
                # Make type() function available for IF conditions
                group_type = DocumentService._get_group_type(group)
                iter_answer[f'type({loop_identifier})'] = group_type
                iter_raw[f'type({loop_identifier})'] = group_type
                for ident in body_ident_names:
                    base_ident = ident.split('.', 1)[0] if '.' in ident else ident
                    if base_ident in same_group and base_ident != loop_identifier:
                        arr = identifier_arrays.get(base_ident)
                        if arr is not None and first_idx < len(arr):
                            fld = ident.split('.', 1)[1] if '.' in ident else None
                            val = DocumentService._format_people_group_field(
                                group, arr, first_idx, fld
                            )
                            iter_answer[ident] = val
                            iter_raw[ident] = val
                instance_body = DocumentService._process_if_blocks(
                    instance_body, iter_answer, iter_raw
                )
                instance_body = DocumentService._process_switch_blocks(
                    instance_body, iter_answer, iter_raw
                )

                # Replace counter tokens AFTER IF processing so only
                # surviving ##% tokens consume counter values.
                instance_body = DocumentService._replace_counter_tokens(instance_body, global_counter)

                output_parts.append(instance_body)

            expanded = ''.join(output_parts)

            # Recursively process nested loops
            expanded = DocumentService._process_peopleloop_blocks(
                expanded, answer_map, raw_map, global_counter, identifier_group_map
            )

            result.append(expanded)
            pos = close_end

        return ''.join(result)

    @staticmethod
    def _resolve_identifier_value(identifier: str, answer_map: dict, raw_answer_map: dict) -> str:
        """Resolve an identifier to its value, supporting dot notation and array indexing.

        For simple identifiers, looks up in answer_map.
        For dot notation (e.g., 'person.relationship'), parses the
        person JSON from raw_answer_map and extracts the field.
        For array indexing (e.g., 'able_to_act[1]'), parses the array and returns the indexed element.
        """
        # Check for array indexing syntax: identifier[N]
        array_match = re.match(r'^([^\[]+)\[(\d+)\](?:\.(.+))?$', identifier)
        if array_match:
            base_identifier = array_match.group(1)
            array_index = int(array_match.group(2)) - 1  # Convert to 0-based
            field_name = array_match.group(3)  # Optional field after array index

            # Get raw value for the base identifier
            raw_json = (raw_answer_map or {}).get(base_identifier, '') or ''
            if raw_json:
                try:
                    parsed = json.loads(raw_json)
                    if isinstance(parsed, list) and 0 <= array_index < len(parsed):
                        item = parsed[array_index]

                        # Decode if double-encoded
                        if isinstance(item, str):
                            try:
                                item = json.loads(item)
                            except (json.JSONDecodeError, TypeError):
                                pass

                        # If field name specified, extract field from dict
                        if field_name and isinstance(item, dict):
                            val = item.get(field_name)
                            return str(val) if val is not None else ''

                        # Otherwise return the whole item
                        if isinstance(item, dict):
                            name = item.get('name', str(item))
                            return DocumentService._extract_plain_name(name)
                        return str(item)
                    elif not isinstance(parsed, list) and array_index == 0:
                        # Scalar value used with [1] index — treat as single-element
                        return str(parsed)
                except (json.JSONDecodeError, TypeError):
                    # Not valid JSON — if index is 0, return the scalar
                    if array_index == 0:
                        return raw_json

        direct = answer_map.get(identifier, '')
        if direct and not DocumentService._is_value_empty(direct):
            return direct

        if '.' in identifier:
            base, field = identifier.split('.', 1)
            raw_json = (raw_answer_map or {}).get(base, '') or ''
            if raw_json:
                try:
                    parsed = json.loads(raw_json)
                    if isinstance(parsed, dict):
                        val = parsed.get(field)
                        if val is not None:
                            return str(val)
                    elif isinstance(parsed, list) and len(parsed) > 0:
                        first = parsed[0]
                        if isinstance(first, str):
                            try:
                                first = json.loads(first)
                            except (json.JSONDecodeError, TypeError):
                                pass
                        if isinstance(first, dict):
                            val = first.get(field)
                            if val is not None:
                                return str(val)
                except (json.JSONDecodeError, TypeError):
                    pass
        return ''

    @staticmethod
    def _evaluate_if_condition(condition_text: str, answer_map: dict, raw_answer_map: dict) -> bool:
        """Evaluate the condition inside {{ IF <condition> }}.

        Supports compound conditions with AND / OR:
          {{IF ident = "a" AND ident2 = "b"}}
          {{IF ident = "a" OR ident2 = "b"}}
          {{IF ident AND ident2 = "b" OR ident3}}  (AND binds tighter than OR)

        Also supports all single-condition forms:
          identifier = "value"     / identifier != "value"
          NOT identifier
          ANY <<identifier>> = "value"   (true if any array element matches)
          NONE <<identifier>> = "value"  (true if no array element matches)
          count(ident) op N / type(ident) = "value"

        Returns True if the content should be included.
        """
        cond = condition_text.strip()

        # --- Compound condition handling (AND / OR) ---
        # Split on OR first (lower precedence), then AND (higher precedence).
        # Use word-boundary split so we don't match inside identifiers or values.
        or_parts = re.split(r'\s+OR\s+', cond, flags=re.IGNORECASE)
        if len(or_parts) > 1:
            # Any OR-group being true makes the whole condition true
            for or_part in or_parts:
                and_parts = re.split(r'\s+AND\s+', or_part.strip(), flags=re.IGNORECASE)
                if all(DocumentService._evaluate_single_condition(p.strip(), answer_map, raw_answer_map) for p in and_parts):
                    return True
            return False

        # Check for AND without OR
        and_parts = re.split(r'\s+AND\s+', cond, flags=re.IGNORECASE)
        if len(and_parts) > 1:
            return all(DocumentService._evaluate_single_condition(p.strip(), answer_map, raw_answer_map) for p in and_parts)

        # Single condition — delegate directly
        return DocumentService._evaluate_single_condition(cond, answer_map, raw_answer_map)

    @staticmethod
    def _evaluate_single_condition(condition_text: str, answer_map: dict, raw_answer_map: dict) -> bool:
        """Evaluate a single (non-compound) condition."""
        cond = condition_text.strip()

        # count() function — returns the length of an array identifier.
        # Supports =, !=, >, <, >=, <= with a numeric right-hand operand.
        _q = r'["\'""''\u00ab\u00bb]'
        count_match = re.match(
            r'count\s*\(\s*(?:<<)?([^>=!\s\}>)]+)(?:>>)?\s*\)\s*(=|!=|>=?|<=?)\s*' + _q + r'?(\d+)' + _q + r'?',
            cond, re.IGNORECASE
        )
        if count_match:
            identifier = count_match.group(1).lower()
            operator = count_match.group(2)
            expected_count = int(count_match.group(3))
            raw_value = (raw_answer_map or {}).get(identifier, '') or (answer_map or {}).get(identifier, '')
            actual_count = 0
            if raw_value:
                try:
                    parsed = json.loads(raw_value)
                    if isinstance(parsed, list):
                        actual_count = len(parsed)
                    elif not DocumentService._is_value_empty(raw_value):
                        actual_count = 1
                except (json.JSONDecodeError, TypeError):
                    if not DocumentService._is_value_empty(raw_value):
                        actual_count = 1
            _logger.debug(f"IF count(): count({identifier}) {operator} {expected_count} -> actual={actual_count}")
            if operator == '=':
                return actual_count == expected_count
            elif operator == '!=':
                return actual_count != expected_count
            elif operator == '>':
                return actual_count > expected_count
            elif operator == '<':
                return actual_count < expected_count
            elif operator == '>=':
                return actual_count >= expected_count
            elif operator == '<=':
                return actual_count <= expected_count
            return False

        # type() function for PEOPLELOOP group classification
        type_match = re.match(
            r'type\s*\(\s*(?:<<)?([^>=!\s\}>)]+)(?:>>)?\s*\)\s*(=|!=)\s*' + _q + r'([^"\'\u201c\u201d\u2018\u2019\u00ab\u00bb]*)' + _q + r'?',
            cond, re.IGNORECASE
        )
        if type_match:
            identifier = type_match.group(1).lower()
            operator = type_match.group(2)
            expected = (type_match.group(3) or '').strip()
            type_key = f'type({identifier})'
            actual = (answer_map or {}).get(type_key, '')
            _logger.debug(f"IF type(): type({identifier}) {operator} '{expected}' -> actual='{actual}'")
            if operator == '!=':
                return actual.lower() != expected.lower()
            return actual.lower() == expected.lower()

        # ANY / NONE aggregate operators for repeatable fields
        any_none_match = re.match(
            r'(ANY|NONE)\s+(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*=\s*' + _q + r'([^"\'\u201c\u201d\u2018\u2019\u00ab\u00bb]*)' + _q + r'?',
            cond, re.IGNORECASE
        )
        if any_none_match:
            quantifier = any_none_match.group(1).upper()
            identifier = any_none_match.group(2).lower()
            expected = any_none_match.group(3) or ''
            raw_value = (raw_answer_map or {}).get(identifier, '') or (answer_map or {}).get(identifier, '')
            # Parse as JSON array; fall back to single-element list
            values: list = []
            if raw_value:
                try:
                    parsed = json.loads(raw_value)
                    if isinstance(parsed, list):
                        for v in parsed:
                            if v is None:
                                values.append('')
                                continue
                            # Decode double-encoded JSON strings
                            item = v
                            if isinstance(item, str):
                                try:
                                    item = json.loads(item)
                                except (json.JSONDecodeError, TypeError):
                                    pass
                            # Extract name from person objects (dicts with 'name' key)
                            if isinstance(item, dict):
                                name = item.get('name', '')
                                plain = DocumentService._extract_plain_name(name) if name else str(item)
                                values.append(plain.lower())
                            else:
                                values.append(str(item).lower())
                    else:
                        values = [str(raw_value).lower()]
                except (json.JSONDecodeError, TypeError, ValueError):
                    values = [str(raw_value).lower()]
            any_match = expected.lower() in values
            return any_match if quantifier == 'ANY' else not any_match

        not_match = re.match(
            r'NOT\s+(?:<<)?([^>=!\s\}>]+)(?:>>)?$', cond, re.IGNORECASE
        )
        if not_match:
            identifier = not_match.group(1).lower()
            value = DocumentService._resolve_identifier_value(identifier, answer_map, raw_answer_map)
            return DocumentService._is_value_empty(value)

        neq_match = re.match(
            r'(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*!=\s*(?:' + _q + r'([^"\'\u201c\u201d\u2018\u2019\u00ab\u00bb]*)' + _q + r'?|(EMPTY|NULL|NONE))',
            cond, re.IGNORECASE
        )
        if neq_match:
            identifier = neq_match.group(1).lower()
            keyword = neq_match.group(3)
            actual = DocumentService._resolve_identifier_value(identifier, answer_map, raw_answer_map)
            if keyword and keyword.upper() in ('EMPTY', 'NULL', 'NONE'):
                return not DocumentService._is_value_empty(actual)
            expected = neq_match.group(2) or ''
            return actual.lower() != expected.lower()

        eq_match = re.match(
            r'(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*=\s*(?:' + _q + r'([^"\'\u201c\u201d\u2018\u2019\u00ab\u00bb]*)' + _q + r'?|(EMPTY|NULL|NONE))',
            cond, re.IGNORECASE
        )
        if eq_match:
            identifier = eq_match.group(1).lower()
            keyword = eq_match.group(3)
            actual = DocumentService._resolve_identifier_value(identifier, answer_map, raw_answer_map)
            if keyword and keyword.upper() in ('EMPTY', 'NULL', 'NONE'):
                return DocumentService._is_value_empty(actual)
            expected = eq_match.group(2) or ''
            return actual.lower() == expected.lower()

        # Identifier-to-identifier comparison (unquoted right side).
        # Supports subscripts and dot notation on both sides, e.g.:
        #   {{IF principal = principal[1]}}
        #   {{IF <<person>> != <<person[2]>>}}
        ident_cmp_match = re.match(
            r'(?:<<)?([^>=!\s\}>]+)(?:>>)?\s*(=|!=)\s*(?:<<)?([^>=!\s\}>]+)(?:>>)?$',
            cond, re.IGNORECASE
        )
        if ident_cmp_match:
            left_ident = ident_cmp_match.group(1).lower()
            operator = ident_cmp_match.group(2)
            right_ident = ident_cmp_match.group(3).lower()
            left_value = DocumentService._resolve_identifier_value(left_ident, answer_map, raw_answer_map)
            right_value = DocumentService._resolve_identifier_value(right_ident, answer_map, raw_answer_map)
            _logger.debug(f"IF ident-cmp: {left_ident}='{left_value}' {operator} {right_ident}='{right_value}'")
            if operator == '!=':
                return left_value.lower() != right_value.lower()
            return left_value.lower() == right_value.lower()

        plain_match = re.match(
            r'(?:<<)?([^>=!\s\}>]+)(?:>>)?$', cond, re.IGNORECASE
        )
        if plain_match:
            identifier = plain_match.group(1).lower()
            value = DocumentService._resolve_identifier_value(identifier, answer_map, raw_answer_map)
            return not DocumentService._is_value_empty(value)

        return True

    @staticmethod
    def _process_if_blocks(text: str, answer_map: dict, raw_answer_map: dict) -> str:
        """Recursively process nested {{ IF }} ... {{ ELSE }} ... {{ END }} blocks.

        Scans *text* left-to-right.  When an {{ IF ... }} tag is found the
        parser counts nesting depth to locate the matching {{ END }} and an
        optional {{ ELSE }} at the same depth.  The if-body (and else-body,
        if present) are recursively processed, then the condition decides
        which branch to keep.
        """
        _if_open_re = DocumentService._IF_OPEN_RE
        _end_re = DocumentService._END_RE
        _else_re = DocumentService._ELSE_RE

        result = []
        pos = 0

        while pos < len(text):
            open_match = _if_open_re.search(text, pos)
            if not open_match:
                result.append(text[pos:])
                break

            result.append(text[pos:open_match.start()])

            condition_text = open_match.group(1)
            body_start = open_match.end()

            depth = 1
            scan = body_start
            body_end = body_start
            else_start = None
            else_end = None
            while depth > 0 and scan < len(text):
                next_open = _if_open_re.search(text, scan)
                next_end = _end_re.search(text, scan)
                next_else = _else_re.search(text, scan)

                if next_end is None:
                    scan = len(text)
                    break

                candidates = [('end', next_end)]
                if next_open:
                    candidates.append(('open', next_open))
                if next_else:
                    candidates.append(('else', next_else))
                candidates.sort(key=lambda c: c[1].start())

                tag_type, tag_match = candidates[0]

                if tag_type == 'open':
                    depth += 1
                    scan = tag_match.end()
                elif tag_type == 'else' and depth == 1:
                    else_start = tag_match.start()
                    else_end = tag_match.end()
                    scan = tag_match.end()
                elif tag_type == 'else':
                    scan = tag_match.end()
                else:
                    depth -= 1
                    if depth == 0:
                        body_end = tag_match.start()
                        scan = tag_match.end()
                    else:
                        scan = tag_match.end()

            if depth != 0:
                result.append(text[open_match.start():])
                break

            if else_start is not None:
                if_body = text[body_start:else_start]
                else_body = text[else_end:body_end]
            else:
                if_body = text[body_start:body_end]
                else_body = None

            if DocumentService._evaluate_if_condition(condition_text, answer_map, raw_answer_map):
                chosen = DocumentService._process_if_blocks(if_body.strip(), answer_map, raw_answer_map)
                if chosen:
                    result.append(chosen)
            elif else_body is not None:
                chosen = DocumentService._process_if_blocks(else_body.strip(), answer_map, raw_answer_map)
                if chosen:
                    result.append(chosen)

            pos = scan

        return ''.join(result)

    @staticmethod
    def _process_switch_blocks(text: str, answer_map: dict, raw_answer_map: dict) -> str:
        """Process {{ SWITCH <<identifier>> }} ... {{ CASE "val" }} ... {{ ELSE }} ... {{ END SWITCH }} blocks.

        Resolves the switch identifier, then walks through CASE branches
        to find the first match (case-insensitive).  If no CASE matches
        and an ELSE branch exists, the ELSE body is used.  The chosen
        body is recursively processed for nested IF / SWITCH blocks.
        """
        _switch_re = DocumentService._SWITCH_OPEN_RE
        _case_re = DocumentService._CASE_RE
        _else_re = DocumentService._ELSE_RE
        _end_switch_re = DocumentService._END_SWITCH_RE

        result: list[str] = []
        pos = 0

        while pos < len(text):
            sw_match = _switch_re.search(text, pos)
            if not sw_match:
                result.append(text[pos:])
                break

            result.append(text[pos:sw_match.start()])

            identifier = sw_match.group(1).strip().lower()
            block_start = sw_match.end()

            # Find matching {{ END SWITCH }}, respecting nested SWITCH blocks
            depth = 1
            scan = block_start
            block_end = None
            while depth > 0 and scan < len(text):
                next_open = _switch_re.search(text, scan)
                next_close = _end_switch_re.search(text, scan)

                if next_close is None:
                    scan = len(text)
                    break

                if next_open and next_open.start() < next_close.start():
                    depth += 1
                    scan = next_open.end()
                else:
                    depth -= 1
                    if depth == 0:
                        block_end = next_close.start()
                        scan = next_close.end()
                    else:
                        scan = next_close.end()

            if depth != 0:
                # Unmatched SWITCH — keep original text
                result.append(text[sw_match.start():])
                break

            block_content = text[block_start:block_end]

            # Resolve the switch identifier value.
            # Support count() function: {{SWITCH count(ident)}}
            count_sw_match = re.match(
                r'count\s*\(\s*(?:[<«‹]{2})?([^>»›\)\s]+)(?:[>»›]{2})?\s*\)',
                identifier, re.IGNORECASE
            )
            if count_sw_match:
                count_ident = count_sw_match.group(1).lower()
                raw_value = (raw_answer_map or {}).get(count_ident, '') or (answer_map or {}).get(count_ident, '')
                actual_count = 0
                if raw_value:
                    try:
                        parsed = json.loads(raw_value)
                        if isinstance(parsed, list):
                            actual_count = len(parsed)
                        elif not DocumentService._is_value_empty(raw_value):
                            actual_count = 1
                    except (json.JSONDecodeError, TypeError):
                        if not DocumentService._is_value_empty(raw_value):
                            actual_count = 1
                switch_value = str(actual_count)
            else:
                switch_value = DocumentService._resolve_identifier_value(
                    identifier, answer_map, raw_answer_map
                ).lower()

            # Parse CASE / ELSE branches from block_content.
            # Collect (tag_start, body_start, case_value_or_None) for each branch.
            branch_tags: list[tuple[int, int, str | None]] = []  # (tag_start, body_start, case_val)

            inner_depth = 0
            inner_pos = 0
            while inner_pos < len(block_content):
                nested_sw = _switch_re.search(block_content, inner_pos)
                nested_end_sw = _end_switch_re.search(block_content, inner_pos)

                if inner_depth > 0:
                    candidates = []
                    if nested_sw:
                        candidates.append(('open', nested_sw))
                    if nested_end_sw:
                        candidates.append(('close', nested_end_sw))
                    if not candidates:
                        break
                    candidates.sort(key=lambda c: c[1].start())
                    tag_type, tag_match = candidates[0]
                    if tag_type == 'open':
                        inner_depth += 1
                    else:
                        inner_depth -= 1
                    inner_pos = tag_match.end()
                    continue

                # At top level — look for CASE, ELSE, or nested SWITCH open
                case_m = _case_re.search(block_content, inner_pos)
                else_m = _else_re.search(block_content, inner_pos)

                candidates = []
                if case_m:
                    candidates.append(('case', case_m))
                if else_m:
                    candidates.append(('else', else_m))
                if nested_sw:
                    candidates.append(('open', nested_sw))

                if not candidates:
                    break

                candidates.sort(key=lambda c: c[1].start())
                tag_type, tag_match = candidates[0]

                if tag_type == 'open':
                    inner_depth += 1
                    inner_pos = tag_match.end()
                elif tag_type == 'case':
                    case_val = tag_match.group(1) if tag_match.group(1) is not None else tag_match.group(2)
                    branch_tags.append((tag_match.start(), tag_match.end(), case_val))
                    inner_pos = tag_match.end()
                elif tag_type == 'else':
                    branch_tags.append((tag_match.start(), tag_match.end(), None))
                    inner_pos = tag_match.end()

            # Build branches: body runs from body_start to the next tag_start (or block end)
            branches: list[tuple[str | None, str]] = []
            for i, (tag_start, body_start, case_val) in enumerate(branch_tags):
                if i + 1 < len(branch_tags):
                    body_text = block_content[body_start:branch_tags[i + 1][0]]
                else:
                    body_text = block_content[body_start:]
                branches.append((case_val, body_text))

            # Select the matching branch
            chosen_body = None
            else_body = None
            for case_val, body in branches:
                if case_val is None:
                    else_body = body
                elif case_val.lower() == switch_value:
                    chosen_body = body
                    break

            if chosen_body is None:
                chosen_body = else_body

            if chosen_body is not None:
                # Recursively process IF and SWITCH blocks in the chosen body
                processed = DocumentService._process_if_blocks(chosen_body.strip(), answer_map, raw_answer_map)
                processed = DocumentService._process_switch_blocks(processed, answer_map, raw_answer_map)
                if processed:
                    result.append(processed)

            pos = scan

        return ''.join(result)

    @staticmethod
    def _process_conditional_sections(text: str, answer_map: dict, raw_answer_map: dict = None) -> str:
        """Process [[ ... ]] conditional sections.

        If any identifier inside is empty, remove the entire section.
        Otherwise, replace identifiers and remove the brackets.
        """
        _raw_map = raw_answer_map or answer_map

        def _process_section(match):
            section_content = match.group(1)

            identifiers_in_section = re.findall(r'<<([^>]+)>>', section_content)

            if not identifiers_in_section:
                return section_content

            for identifier in identifiers_in_section:
                value = DocumentService._resolve_identifier_value(
                    identifier.lower(), answer_map, _raw_map
                )
                if DocumentService._is_value_empty(value):
                    _logger.debug(f"Identifier '{identifier}' is empty, removing conditional section")
                    return ''

            result = section_content
            for identifier in identifiers_in_section:
                value = DocumentService._resolve_identifier_value(
                    identifier.lower(), answer_map, _raw_map
                )
                result = result.replace(f'<<{identifier}>>', value)
            return result

        return DocumentService._CONDITIONAL_RE.sub(_process_section, text)

    @staticmethod
    def _replace_counter_tokens(text: str, global_counter: list) -> str:
        """Replace ##, ###, ##%, ##A tokens with running counter values.
        Also handle #/A to reset the alphabetical counter.

        Args:
            text: Text containing counter tokens
            global_counter: Mutable [int] shared counter state

        Returns:
            Text with counter tokens replaced
        """
        # Process text sequentially to handle resets and counters in order
        # Combine both patterns to process in sequence
        combined_pattern = re.compile(r'(##/|#/A)|(###|##%|##A|##V|##)(?:\+(\d*))?')

        def _replacer(match):
            if match.group(1):  # Reset token ##/ or #/A
                global_counter[0] = 0
                return ''
            else:  # Counter token
                token = match.group(2)
                plus_str = match.group(3)
                inc = int(plus_str) if plus_str else 1
                global_counter[0] += inc
                return DocumentService._counter_to_str(token, global_counter[0])

        return combined_pattern.sub(_replacer, text)

    @staticmethod
    def _replace_identifiers(text: str, answer_map: dict, raw_answer_map: dict,
                             conjunction_map: dict, identifier_group_map: dict) -> str:
        """Replace all <<identifier>> tokens with their values.

        Handles:
        - Array indexing: <<identifier[0]>> for first item
        - Dot notation for person fields: <<person.name>>
        - Combined: <<person[0].name>> for first person's name
        - JSON arrays with conjunction joining
        - Scalar values
        """
        _conj_map = conjunction_map or {}
        _id_grp_map = identifier_group_map or {}
        _raw_map = raw_answer_map or {}

        def _replace(match):
            """Wrapper that strips HTML tags from the substituted value."""
            return DocumentService._strip_value_html(_replace_raw(match))

        def _replace_raw(match):
            full_match = match.group(0)
            identifier = match.group(1).lower()
            array_index_str = match.group(2)  # Array index like [1], [2], etc.
            field_name = match.group(3)  # Field name after dot or array index

            _logger.debug(f"Replacing identifier: {full_match} -> identifier={identifier}, in_map={identifier in answer_map}")

            # count() function: <<count(identifier)>> outputs the array length
            count_fn_match = re.match(r'^count\((.+)\)$', identifier, re.IGNORECASE)
            if count_fn_match:
                count_ident = count_fn_match.group(1).lower()
                raw_value = _raw_map.get(count_ident, '') or answer_map.get(count_ident, '')
                if raw_value:
                    try:
                        parsed = json.loads(raw_value)
                        if isinstance(parsed, list):
                            return str(len(parsed))
                        elif not DocumentService._is_value_empty(raw_value):
                            return '1'
                    except (json.JSONDecodeError, TypeError):
                        if not DocumentService._is_value_empty(raw_value):
                            return '1'
                return '0'

            # Convert array index to integer (1-based, so subtract 1 for 0-based array access)
            array_index = int(array_index_str) - 1 if array_index_str else None

            # Check for 2D array indexing: identifier[1][2]
            # This handles repeatable questions inside repeatable parents (e.g., conditional followups)
            nested_array_match = None
            if field_name and re.match(r'^\[(\d+)\]', field_name):
                nested_array_match = re.match(r'^\[(\d+)\](?:\.(.+))?$', field_name)
                if nested_array_match:
                    second_index_str = nested_array_match.group(1)
                    second_index = int(second_index_str) - 1  # Convert to 0-based
                    remaining_field = nested_array_match.group(2)  # Any field after second index

                    # Access 2D array: data[array_index][second_index]
                    raw_json = _raw_map.get(identifier, '') or ''
                    if raw_json:
                        try:
                            data = json.loads(raw_json)
                            if isinstance(data, list) and 0 <= array_index < len(data):
                                inner_array = data[array_index]
                                # Decode if double-encoded
                                if isinstance(inner_array, str):
                                    try:
                                        inner_array = json.loads(inner_array)
                                    except (json.JSONDecodeError, TypeError):
                                        pass

                                if isinstance(inner_array, list) and 0 <= second_index < len(inner_array):
                                    item = inner_array[second_index]

                                    # Decode if double-encoded
                                    if isinstance(item, str):
                                        try:
                                            item = json.loads(item)
                                        except (json.JSONDecodeError, TypeError):
                                            pass

                                    # If there's a field after the second index, extract it
                                    if remaining_field and isinstance(item, dict):
                                        field_value = item.get(remaining_field)
                                        return str(field_value) if field_value is not None else ''

                                    # Otherwise return the item
                                    if isinstance(item, dict):
                                        return item.get('name', str(item))
                                    return str(item)
                        except (json.JSONDecodeError, TypeError):
                            pass
                    return ''

            # Handle array indexing with optional field access
            # e.g., <<identifier[1]>>, <<identifier[1].name>>
            if array_index is not None:
                raw_json = _raw_map.get(identifier, '') or ''
                formatted = answer_map.get(identifier, '')

                # 1) Try raw values first (preserves JSON arrays for proper indexing)
                if raw_json:
                    try:
                        data = json.loads(raw_json)
                        if isinstance(data, list):
                            if 0 <= array_index < len(data):
                                item = data[array_index]

                                # Decode if double-encoded
                                decoded = DocumentService._decode_json_item(item)

                                # If field name specified, extract field from dict
                                if field_name and isinstance(decoded, dict):
                                    field_value = decoded.get(field_name)
                                    return str(field_value) if field_value is not None else ''

                                # For dates, format nicely
                                if not field_name and isinstance(decoded, str):
                                    try:
                                        dt = datetime.strptime(decoded, '%Y-%m-%d')
                                        return dt.strftime('%B %-d, %Y')
                                    except (ValueError, TypeError):
                                        pass

                                # Otherwise return the whole item
                                if isinstance(decoded, dict):
                                    name = decoded.get('name', str(decoded))
                                    name = DocumentService._extract_plain_name(name)
                                    return name
                                return str(decoded)
                            else:
                                # Index out of range
                                return ''
                    except (json.JSONDecodeError, TypeError):
                        pass

                # 2) Try formatted values as JSON array
                if not field_name and formatted:
                    try:
                        data = json.loads(formatted)
                        if isinstance(data, list) and 0 <= array_index < len(data):
                            item = data[array_index]
                            if isinstance(item, dict):
                                name = item.get('name', str(item))
                                name = DocumentService._extract_plain_name(name)
                                return name
                            return str(item)
                    except (json.JSONDecodeError, TypeError):
                        pass

                # 3) Scalar fallback: value is not a JSON array but index is 0
                if array_index == 0:
                    if raw_json:
                        # Format dates
                        try:
                            dt = datetime.strptime(raw_json, '%Y-%m-%d')
                            return dt.strftime('%B %-d, %Y')
                        except (ValueError, TypeError):
                            pass
                        return raw_json
                    if formatted:
                        return formatted
                return ''

            # Handle dot notation without array index (legacy behavior)
            # e.g., <<person.name>>
            if field_name:
                raw_json = _raw_map.get(identifier, '') or ''
                formatted = answer_map.get(identifier, '')

                if raw_json:
                    try:
                        person_data = json.loads(raw_json)

                        if isinstance(person_data, dict):
                            field_value = person_data.get(field_name)
                            if field_value is not None:
                                return str(field_value)
                        elif isinstance(person_data, list) and len(person_data) > 0:
                            if field_name == 'name':
                                return formatted if formatted else ''
                            first = person_data[0]
                            if isinstance(first, str):
                                try:
                                    first = json.loads(first)
                                except (json.JSONDecodeError, TypeError):
                                    pass
                            if isinstance(first, dict):
                                field_value = first.get(field_name)
                                if field_value is not None:
                                    return str(field_value)
                    except (json.JSONDecodeError, TypeError):
                        pass

                if formatted:
                    try:
                        person_data = json.loads(formatted)
                        if isinstance(person_data, dict):
                            field_value = person_data.get(field_name)
                            if field_value is not None:
                                return str(field_value)
                    except (json.JSONDecodeError, TypeError):
                        if field_name == 'name':
                            return formatted

                return ''

            # Handle simple identifier without index or field (legacy behavior)
            value = answer_map.get(identifier, '')
            if not DocumentService._is_value_empty(value):
                try:
                    parsed = json.loads(value)
                    if isinstance(parsed, list) and len(parsed) > 0:
                        items = []
                        for item in parsed:
                            decoded = DocumentService._decode_json_item(item)
                            if isinstance(decoded, dict):
                                item_str = DocumentService._extract_plain_name(decoded.get('name', str(decoded)))
                            else:
                                item_str = str(decoded)
                            items.append(item_str)

                        group_id = _id_grp_map.get(identifier)
                        conjunctions = _conj_map.get(group_id, []) if group_id else []

                        if len(items) == 1:
                            return items[0]

                        result_parts = [items[0]]
                        for i in range(1, len(items)):
                            conj = conjunctions[i] if i < len(conjunctions) else 'and'
                            if not conj:
                                conj = 'and'
                            if conj.lower() == 'then':
                                result_parts.append(f', then {items[i]}')
                            elif i == len(items) - 1 and conj.lower() in ('and', 'or'):
                                if len(items) > 2:
                                    result_parts.append(f', {conj} {items[i]}')
                                else:
                                    result_parts.append(f' {conj} {items[i]}')
                            else:
                                result_parts.append(f', {items[i]}')
                        return ''.join(result_parts)
                except (json.JSONDecodeError, TypeError):
                    pass
                return value
            return ''

        # Avoid replacing identifiers inside {{ }} control-flow tags
        _tag_re = re.compile(r'(\{\{.*?\}\})')
        parts = _tag_re.split(text)
        for i in range(len(parts)):
            if not parts[i].startswith('{{'):
                parts[i] = DocumentService._IDENTIFIER_RE.sub(_replace, parts[i])
        return ''.join(parts)

    @staticmethod
    def _process_macros(template_content: str) -> str:
        """
        Process macro definitions and usages in template content.

        Macro syntax:
        - Definition: @@ macro_name @@ macro content with <<identifiers>> @@
        - Usage: @ macro_name @

        Definitions are extracted and removed from the template.
        Usages are replaced with the macro definition content.

        Args:
            template_content: Template content with macro definitions and usages

        Returns:
            Template content with macros expanded and definitions removed
        """
        # Regex to match macro definitions: @@ name @@ content @@
        # Captures: (1) macro name, (2) macro content
        macro_def_pattern = r'@@\s*(\w+)\s*@@\s*(.*?)\s*@@'

        # Extract all macro definitions
        macros = {}
        for match in re.finditer(macro_def_pattern, template_content, re.DOTALL):
            macro_name = match.group(1).strip()
            macro_content = match.group(2).strip()
            macros[macro_name] = macro_content
            _logger.info(f"Defined macro: {macro_name}")

        # Remove all macro definitions from template
        template_content = re.sub(macro_def_pattern, '', template_content, flags=re.DOTALL)

        # Replace macro usages with their definitions
        # Regex to match macro usage: @ name @
        macro_usage_pattern = r'@\s*(\w+)\s*@'

        def replace_macro(match):
            macro_name = match.group(1).strip()
            if macro_name in macros:
                _logger.info(f"Expanding macro: {macro_name}")
                return macros[macro_name]
            else:
                _logger.warning(f"Macro not found: {macro_name}")
                return match.group(0)  # Return original if macro not defined

        template_content = re.sub(macro_usage_pattern, replace_macro, template_content)

        return template_content.strip()

    @staticmethod
    def _process_formatting_tags(text: str) -> str:
        """
        Process formatting tags for alignment, tabs, and indentation.

        Supported tags:
        - <center>text</center> or <CENTER>text</CENTER> - Center alignment
        - <right>text</right> or <RIGHT>text</RIGHT> - Right alignment
        - <indent>text</indent> or <INDENT>text</INDENT> - Indent text
        - <tab> or <TAB> - Insert a tab character

        These tags are converted to HTML that the HTMLToWordConverter can process.
        If there are <cr> tags inside formatting tags, lines are joined with <br/> in one
        paragraph so Word does not add extra space between lines (vs one <p> per line).
        """
        # Convert <tab> to actual tab character
        text = re.sub(r'<[Tt][Aa][Bb]>', '\t', text)

        def apply_alignment(match, alignment_class):
            """Helper to apply alignment to content, handling <cr> tokens inside."""
            content = match.group(1)
            # Split by <cr> or <CR> tokens; single <p> with <br/> avoids paragraph spacing
            lines = [line.strip() for line in re.split(r'<[Cc][Rr]>', content) if line.strip()]
            if not lines:
                return ''
            inner = '<br/>'.join(lines)
            # Single <p> per block — do not emit </p>...<p> glue (orphan <p> + <cr> caused huge Word gaps).
            return f'<p {alignment_class}>{inner}</p>'

        # Convert <center>...</center> to HTML with alignment class
        # Handle <cr> inside by applying centering to each line
        text = re.sub(
            r'<[Cc][Ee][Nn][Tt][Ee][Rr]>(.*?)</[Cc][Ee][Nn][Tt][Ee][Rr]>',
            lambda m: apply_alignment(m, 'class="ql-align-center"'),
            text,
            flags=re.DOTALL
        )

        # Convert <right>...</right> to HTML with alignment class
        text = re.sub(
            r'<[Rr][Ii][Gg][Hh][Tt]>(.*?)</[Rr][Ii][Gg][Hh][Tt]>',
            lambda m: apply_alignment(m, 'class="ql-align-right"'),
            text,
            flags=re.DOTALL
        )

        # Convert <indent>...</indent> to HTML with left margin/indent
        text = re.sub(
            r'<[Ii][Nn][Dd][Ee][Nn][Tt]>(.*?)</[Ii][Nn][Dd][Ee][Nn][Tt]>',
            lambda m: apply_alignment(m, 'style="margin-left: 48px;"'),
            text,
            flags=re.DOTALL
        )

        return text

    # Sentinels (PUA + delimiter) so merge passes treat bold/italic/etc. as opaque text.
    # Innermost tags are replaced first; nested same-type tags are unwrapped over iterations.
    _QF_STRONG_O = '\ue000S1\ue001'
    _QF_STRONG_C = '\ue000S0\ue001'
    _QF_BOLD_O = '\ue000B1\ue001'
    _QF_BOLD_C = '\ue000B0\ue001'
    _QF_EM_O = '\ue000E1\ue001'
    _QF_EM_C = '\ue000E0\ue001'
    _QF_I_O = '\ue000I1\ue001'
    _QF_I_C = '\ue000I0\ue001'
    _QF_U_O = '\ue000U1\ue001'
    _QF_U_C = '\ue000U0\ue001'

    @staticmethod
    def _protect_quill_inline_format(html: str) -> str:
        """Replace Quill inline tags with sentinels so merge logic does not drop formatting."""
        patterns = [
            (r'<u\b[^>]*>((?:(?!<u\b).)*?)</u>', DocumentService._QF_U_O, DocumentService._QF_U_C),
            (r'<em\b[^>]*>((?:(?!<em\b).)*?)</em>', DocumentService._QF_EM_O, DocumentService._QF_EM_C),
            (r'<i\b[^>]*>((?:(?!<i\b).)*?)</i>', DocumentService._QF_I_O, DocumentService._QF_I_C),
            (r'<strong\b[^>]*>((?:(?!<strong\b).)*?)</strong>', DocumentService._QF_STRONG_O, DocumentService._QF_STRONG_C),
            (r'<b\b[^>]*>((?:(?!<b\b).)*?)</b>', DocumentService._QF_BOLD_O, DocumentService._QF_BOLD_C),
        ]
        for _ in range(64):
            changed = False
            for open_re, open_tok, close_tok in patterns:
                def _repl(m, o=open_tok, c=close_tok):
                    return o + m.group(1) + c

                new = re.sub(open_re, _repl, html, flags=re.DOTALL | re.IGNORECASE)
                if new != html:
                    changed = True
                    html = new
            if not changed:
                break
        return html

    @staticmethod
    def _restore_quill_inline_format(text: str) -> str:
        """Restore Quill inline formatting sentinels to HTML for HTMLToWordConverter."""
        pairs = [
            (DocumentService._QF_STRONG_O, '<strong>'),
            (DocumentService._QF_STRONG_C, '</strong>'),
            (DocumentService._QF_BOLD_O, '<b>'),
            (DocumentService._QF_BOLD_C, '</b>'),
            (DocumentService._QF_EM_O, '<em>'),
            (DocumentService._QF_EM_C, '</em>'),
            (DocumentService._QF_I_O, '<i>'),
            (DocumentService._QF_I_C, '</i>'),
            (DocumentService._QF_U_O, '<u>'),
            (DocumentService._QF_U_C, '</u>'),
        ]
        for tok, tag in pairs:
            text = text.replace(tok, tag)
        return text

    @staticmethod
    def _unwrap_inline_tags_splitting_template_syntax(html: str) -> str:
        """
        Unwrap strong/em/b/i/u that only wrap delimiter fragments or whitespace.

        Quill can split @macro@ or <<id>> across tags; those must merge as plain text
        before identifier/macro passes. Formatting that wraps real words is left for
        _protect_quill_inline_format.
        """
        for _ in range(12):
            prev = html
            # Whitespace-only (or empty) — keeps the whitespace run
            html = re.sub(
                r'<(strong|em|b|i|u)\b[^>]*>(\s*)</\1>',
                lambda m: m.group(2),
                html,
                flags=re.IGNORECASE,
            )
            for inner in ('@', '<<', '>>'):
                esc = re.escape(inner)
                html = re.sub(
                    rf'<(strong|em|b|i|u)\b[^>]*>{esc}</\1>',
                    inner,
                    html,
                    flags=re.IGNORECASE,
                )
            if html == prev:
                break
        return html

    @staticmethod
    def _merge_consecutive_same_class_paragraphs(html: str, class_token: str) -> str:
        """Merge adjacent <p> tags that share a Quill alignment class into one <p> with <br/>."""
        esc = re.escape(class_token)
        pat = rf'(<p\b[^>]*\b{esc}\b[^>]*>)(.*?)(</p>)(\s*)(<p\b[^>]*\b{esc}\b[^>]*>)(.*?)(</p>)'
        for _ in range(64):
            new = re.sub(pat, r'\1\2<br/>\6\3', html, flags=re.DOTALL | re.IGNORECASE)
            if new == html:
                break
            html = new
        return html

    @staticmethod
    def _merge_stacked_quill_aligned_paragraphs(html: str) -> str:
        """
        Separate <center> blocks become separate <p class="ql-align-*">; Word adds space between
        paragraphs. Merge consecutive same-alignment blocks so title lines stack tightly.
        """
        for token in ('ql-align-center', 'ql-align-right', 'ql-align-justify'):
            html = DocumentService._merge_consecutive_same_class_paragraphs(html, token)
        return html

    @staticmethod
    def _merge_template(template_content: str, answer_map: dict, raw_answer_map: dict = None, conjunction_map: dict = None, identifier_group_map: dict = None) -> str:
        """
        Merge template content with answer values.

        Orchestrates seven passes in order:
        0. Macros — extract and expand macro definitions
        1. FOR EACH loops — expand repeatable blocks
        2. IF / ELSE conditionals — evaluate nested conditional blocks
        3. [[ ... ]] conditional sections — remove sections with empty identifiers
        4. Counter tokens (##, ###, ##%) — replace with running numbers
        5. Identifier replacement — replace remaining <<identifier>> tokens
        6. Formatting tags — process alignment, tabs, and indentation tags

        Args:
            template_content: Template markdown content
            answer_map: Dictionary mapping identifiers to formatted answer values
            raw_answer_map: Dictionary mapping identifiers to raw (unformatted) values
            conjunction_map: {repeatable_group_id: [conjunctions]} for joining arrays
            identifier_group_map: {identifier: repeatable_group_id} for repeatable questions

        Returns:
            Merged content with identifiers replaced
        """
        import html

        # Strip BOM (byte order mark) — Python's strip() does NOT remove U+FEFF
        template_content = template_content.replace('\ufeff', '')

        # Pre-process: strip <span> only so @macro@ and <<identifier>> are not split across tags.
        # Quill uses spans for colors/sizes; semantic tags (strong, em, …) are preserved via sentinels.
        # Use .*? because template syntax like <<id>> contains angle brackets.
        for _ in range(5):
            cleaned = re.sub(r'<span\b[^>]*>(.*?)</span>', r'\1', template_content, flags=re.DOTALL | re.IGNORECASE)
            if cleaned == template_content:
                break
            template_content = cleaned

        # Protect page break tokens before unescaping (user-typed <p> is stored as &lt;p&gt;)
        template_content = re.sub(r'&lt;[pP]&gt;', '__PAGE_BREAK__', template_content)

        # Decode HTML entities like &lt; &gt; &amp;
        template_content = html.unescape(template_content)

        # Unwrap delimiter-only inline tags before protecting bold/italic (so @ / << / >> still match).
        template_content = DocumentService._unwrap_inline_tags_splitting_template_syntax(template_content)

        template_content = DocumentService._protect_quill_inline_format(template_content)

        # Flatten: strip HTML paragraph/break tags and literal newlines.
        # The editor uses line breaks for readability only; use <cr> or <CR> for line breaks
        # in the merged document (rendered as <br/>, not new paragraphs, to avoid extra gap).
        template_content = re.sub(r'</p>\s*<p(?:\s[^>]*)?\s*>', ' ', template_content)
        template_content = re.sub(r'</?p(?:\s[^>]*)?\s*>', '', template_content)
        template_content = re.sub(r'<br(?:\s[^>]*)?\s*/?>', ' ', template_content)
        template_content = template_content.replace('\n', ' ').replace('\r', '')
        template_content = re.sub(r'  +', ' ', template_content).strip()

        raw_map = raw_answer_map if raw_answer_map else answer_map
        _id_grp_map = identifier_group_map or {}
        global_counter = [0]

        # Pass 0: Process macros
        merged = DocumentService._process_macros(template_content)

        # Pass 1: FOR EACH loops
        merged = DocumentService._process_foreach_blocks(
            merged, answer_map, raw_map, global_counter, _id_grp_map
        )

        # Pass 1.5: PEOPLELOOP blocks (groups person entries by 'then' conjunction)
        merged = DocumentService._process_peopleloop_blocks(
            merged, answer_map, raw_map, global_counter, _id_grp_map
        )

        # Pass 2: IF / ELSE conditionals
        merged = DocumentService._process_if_blocks(merged, answer_map, raw_answer_map)

        # Pass 2.5: SWITCH / CASE blocks
        merged = DocumentService._process_switch_blocks(merged, answer_map, raw_answer_map)

        # Pass 3: [[ ... ]] conditional sections
        merged = DocumentService._process_conditional_sections(merged, answer_map, raw_answer_map)

        # Pass 4: Counter tokens
        merged = DocumentService._replace_counter_tokens(merged, global_counter)

        # Pass 5: Identifier replacement
        merged = DocumentService._replace_identifiers(
            merged, answer_map, raw_answer_map, conjunction_map, identifier_group_map
        )

        # Pass 6: Formatting tags (alignment, tabs, indentation)
        merged = DocumentService._process_formatting_tags(merged)

        # <cr> after </p> or before <p> is redundant (block boundary already breaks lines); if turned
        # into <br/> it becomes </p><br/>… which Word renders as an extra paragraph + break (large gap).
        merged = re.sub(r'</p>\s*<cr>\s*', '</p>', merged, flags=re.IGNORECASE)
        merged = re.sub(r'<cr>\s*(?=<p\b)', '', merged, flags=re.IGNORECASE)
        merged = DocumentService._merge_stacked_quill_aligned_paragraphs(merged)
        # Remaining <cr> = line breaks inside one block
        merged = re.sub(r'\s*<[Cc][Rr]>\s*', '<br/>', merged)

        # Replace page break placeholders with page break marker
        merged = merged.replace('__PAGE_BREAK__', '</p><pagebreak/><p>')

        merged = DocumentService._restore_quill_inline_format(merged)

        # Drop stray </p><br/> before plain text or before the next <p> (fixes oversized gaps in Word).
        merged = re.sub(r'</p>\s*<br\s*/?>\s*(?=<p\b)', '</p>', merged, flags=re.IGNORECASE)
        merged = re.sub(r'</p>\s*<br\s*/?>\s*(?=[^<\s])', '</p>', merged, flags=re.IGNORECASE)

        # Wrap in <p> only when there is no leading <p> (avoid nested <p> and duplicate paragraph gaps).
        _m = merged.strip()
        if _m and not re.match(r'^<p\b', _m, re.IGNORECASE):
            merged = f'<p>{merged}</p>'

        # HTML-aware cleanup: remove empty <p> tags left behind by removed blocks
        # Matches <p ...> containing only whitespace, &nbsp;, and/or <br> tags
        _empty_p = r'<p[^>]*>\s*(?:&nbsp;|\s|<br\s*/?>)*\s*</p>'
        merged = re.sub(_empty_p, '', merged, flags=re.IGNORECASE)

        # Preserve multiple spaces for HTML rendering by converting runs of 2+
        # regular/non-breaking spaces into &nbsp; entities
        merged = re.sub(r'[\x20\xa0]{2,}', lambda m: '&nbsp;' * len(m.group(0)), merged)

        # Collapse 3+ consecutive newlines (with optional whitespace) down to 2
        merged = re.sub(r'(\s*\n){3,}', '\n\n', merged)

        # Strip leading/trailing whitespace from the final output
        merged = merged.strip()

        return merged

    @staticmethod
    def preview_document(
        db: Session,
        session_id: int,
        template_id: int,
        user_id: int
    ) -> dict:
        """
        Preview a document merge without saving.

        Args:
            db: Database session
            session_id: Session ID
            template_id: Template ID
            user_id: User ID

        Returns:
            Preview data including merged content and missing identifiers
        """
        # Get template
        template = db.query(Template).filter(
            Template.id == template_id,
            Template.is_active == True
        ).first()

        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )

        # Get session
        session = db.query(InputForm).filter(
            InputForm.id == session_id,
            InputForm.user_id == user_id
        ).first()

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found"
            )

        # Get all answers with their questions in a single joined query
        answer_pairs = db.query(SessionAnswer, Question).join(
            Question, SessionAnswer.question_id == Question.id
        ).filter(
            SessionAnswer.session_id == session_id
        ).all()

        answer_map = DocumentService._build_answer_map(answer_pairs)

        # Build raw answer map (unformatted) so FOREACH can parse JSON arrays
        raw_answer_map = DocumentService._build_raw_answer_map(answer_pairs)

        # Build conjunction info for repeatable groups
        conj_map, id_grp_map = DocumentService._build_conjunction_info(answer_pairs)

        # Get template identifiers
        template_identifiers = template.extract_identifiers()

        # Find missing identifiers
        missing_identifiers = [
            identifier for identifier in template_identifiers
            if identifier.lower() not in answer_map
        ]

        # Merge content
        merged_content = DocumentService._merge_template(
            template.markdown_content,
            answer_map,
            raw_answer_map,
            conj_map,
            id_grp_map
        )

        return {
            "template_name": template.name,
            "session_client": session.client_identifier,
            "markdown_content": merged_content,
            "missing_identifiers": missing_identifiers,
            "available_identifiers": list(answer_map.keys())
        }

    @staticmethod
    def get_document(
        db: Session,
        document_id: int,
        user_id: int
    ) -> Optional[GeneratedDocument]:
        """
        Get a generated document by ID.

        Args:
            db: Database session
            document_id: Document ID
            user_id: User ID

        Returns:
            Generated document if found and user has access
        """
        return db.query(GeneratedDocument).join(
            InputForm,
            GeneratedDocument.session_id == InputForm.id
        ).filter(
            GeneratedDocument.id == document_id,
            InputForm.user_id == user_id
        ).first()

    @staticmethod
    def list_documents(
        db: Session,
        user_id: int,
        skip: int = 0,
        limit: int = 100
    ) -> Tuple[List[GeneratedDocument], int]:
        """
        List generated documents for a user.

        Args:
            db: Database session
            user_id: User ID
            skip: Number of records to skip
            limit: Maximum number of records to return

        Returns:
            Tuple of (documents list, total count)
        """
        query = db.query(GeneratedDocument).join(
            InputForm,
            GeneratedDocument.session_id == InputForm.id
        ).filter(
            InputForm.user_id == user_id
        )

        total = query.count()
        documents = query.order_by(GeneratedDocument.generated_at.desc()).offset(skip).limit(limit).all()

        # Download markdown content from S3 for each document
        for doc in documents:
            if doc.s3_key:
                try:
                    doc.markdown_content = s3_service.download_markdown(doc.s3_key)
                except Exception as e:
                    _logger.error(f"Failed to download markdown from S3 for document {doc.id}: {e}")
                    doc.markdown_content = "[Error loading document content]"

        return documents, total

    @staticmethod
    def delete_document(
        db: Session,
        document_id: int,
        user_id: int
    ) -> bool:
        """
        Delete a generated document.

        Args:
            db: Database session
            document_id: Document ID
            user_id: User ID

        Returns:
            True if deleted, False if not found
        """
        document = DocumentService.get_document(db, document_id, user_id)
        if not document:
            return False

        # Delete from S3 if s3_key exists
        if document.s3_key:
            try:
                s3_service.delete_document(document.s3_key)
            except Exception as e:
                _logger.error(f"Failed to delete document from S3: {e}")
                # Continue with database deletion even if S3 deletion fails

        db.delete(document)
        db.commit()

        return True

    @staticmethod
    def merge_document(
        db: Session,
        session_id: int,
        template_id: int,
        user_id: int
    ) -> bytes:
        """
        Merge a template with session data and return a Word document.

        Args:
            db: Database session
            session_id: Document session ID
            template_id: Template ID
            user_id: User ID

        Returns:
            Bytes of the generated Word document
        """
        # Get template
        template = db.query(Template).filter(
            Template.id == template_id,
            Template.is_active == True
        ).first()

        if not template:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )

        # Get session (verify user owns it)
        session = db.query(InputForm).filter(
            InputForm.id == session_id,
            InputForm.user_id == user_id
        ).first()

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found"
            )

        # Get all answers with their questions in a single joined query
        answer_pairs = db.query(SessionAnswer, Question).join(
            Question, SessionAnswer.question_id == Question.id
        ).filter(
            SessionAnswer.session_id == session_id
        ).all()

        # Build answer maps and conjunction info from the joined pairs
        raw_answer_map = DocumentService._build_raw_answer_map(answer_pairs)
        answer_map = DocumentService._build_answer_map(answer_pairs)
        conj_map, id_grp_map = DocumentService._build_conjunction_info(answer_pairs)

        # Get template markdown content and merge using the shared _merge_template function
        # This handles all conditional logic ([[ ]], {{ IF }}, etc.) and identifier replacement
        content = template.markdown_content or ""
        merged_content = DocumentService._merge_template(content, answer_map, raw_answer_map, conj_map, id_grp_map)

        # Create a Word document
        doc = Document()

        # Clean and prepare HTML content
        # The merged_content now contains HTML from the rich text editor
        html_content = merged_content

        # Write HTML to debug file for inspection
        try:
            with open('/tmp/quill_html_debug.html', 'w') as f:
                f.write("=== BEFORE PROCESSING ===\n")
                f.write(html_content)
                f.write("\n\n")
        except:
            pass

        # Remove Quill editor wrapper divs if present
        html_content = re.sub(r'<div class="ql-editor[^"]*"[^>]*>', '', html_content)
        html_content = html_content.replace('</div>', '')

        # Normalize line breaks - Quill uses <p> tags, ensure we don't double them
        # Remove empty paragraphs that cause double spacing
        html_content = re.sub(r'<p>\s*<br\s*/?>\s*</p>', '<p></p>', html_content)
        html_content = re.sub(r'<p>\s*</p>', '', html_content)

        # Do NOT convert <br/> to </p><p>. _merge_template maps <cr> to <br/> for tight line
        # breaks; HTMLToWordConverter turns <br> into Word line breaks within the paragraph.
        # Replacing <br/> with new <p> blocks reintroduces paragraph spacing (large gaps in Word).

        # Ensure content is wrapped in paragraphs
        if not html_content.strip().startswith('<p'):
            html_content = f'<p>{html_content}</p>'

        # Write processed HTML to debug file
        try:
            with open('/tmp/quill_html_debug.html', 'a') as f:
                f.write("=== AFTER PROCESSING ===\n")
                f.write(html_content)
        except:
            pass

        # Parse HTML and convert to Word with custom parser
        parser = HTMLToWordConverter(doc)
        parser.feed(html_content)

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        return doc_bytes.getvalue()
