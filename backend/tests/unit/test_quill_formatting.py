"""Test Quill HTML formatting in Word output - tabs and alignment."""

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.services.document_service import HTMLToWordConverter
import io


class TestQuillAlignment:
    """Test that Quill alignment formatting appears in Word output."""

    def test_center_alignment_class_based(self):
        """Quill center alignment using ql-align-center class."""
        html = '<p class="ql-align-center">Centered text</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_right_alignment_class_based(self):
        """Quill right alignment using ql-align-right class."""
        html = '<p class="ql-align-right">Right aligned text</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_justify_alignment_class_based(self):
        """Quill justify alignment using ql-align-justify class."""
        html = '<p class="ql-align-justify">Justified text</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_center_alignment_style_based(self):
        """Center alignment using inline style."""
        html = '<p style="text-align: center;">Centered via style</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_mixed_alignment_formats(self):
        """Multiple paragraphs with different alignments."""
        html = '''
        <p>Left aligned (default)</p>
        <p class="ql-align-center">Center aligned</p>
        <p style="text-align: right;">Right aligned</p>
        '''
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) >= 3
        # First paragraph should be left (default/None)
        assert doc.paragraphs[0].alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)
        # Second should be center
        assert doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
        # Third should be right
        assert doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.RIGHT


class TestQuillTabs:
    """Test that tab characters from Quill appear in Word output."""

    def test_single_tab_character(self):
        """Tab character should be preserved in Word output."""
        html = '<p>Before\tAfter</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        # Check that the paragraph contains runs with tab character
        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert '\t' in text

    def test_multiple_tabs(self):
        """Multiple tabs should all be preserved."""
        html = '<p>Col1\tCol2\tCol3</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert text.count('\t') == 2

    def test_tab_with_formatting(self):
        """Tabs should work with other formatting like bold."""
        html = '<p><strong>Bold</strong>\tNormal</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        runs = doc.paragraphs[0].runs
        # Should have at least 3 runs: bold text, tab, normal text
        assert len(runs) >= 3
        # Check that tab is present
        text = ''.join(run.text for run in runs)
        assert '\t' in text

    def test_tab_at_paragraph_start(self):
        """Tab at the beginning of a paragraph."""
        html = '<p>\tIndented text</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert text.startswith('\t')


class TestQuillCombined:
    """Test combined formatting scenarios."""

    def test_centered_text_with_tabs(self):
        """Center aligned paragraph containing tabs."""
        html = '<p class="ql-align-center">Name\tAddress\tPhone</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        para = doc.paragraphs[0]

        # Check alignment
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

        # Check tabs are preserved
        text = ''.join(run.text for run in para.runs)
        assert text.count('\t') == 2

    def test_right_aligned_with_tab(self):
        """Right aligned text with tab character."""
        html = '<p style="text-align: right;">Label:\tValue</p>'
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) > 0
        para = doc.paragraphs[0]

        assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        text = ''.join(run.text for run in para.runs)
        assert '\t' in text
