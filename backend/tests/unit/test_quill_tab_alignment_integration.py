"""Integration test: Verify Quill tabs and alignment work in full document generation flow."""

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.services.document_service import HTMLToWordConverter


class TestQuillTabAlignmentIntegration:
    """Test realistic Quill HTML scenarios with tabs and alignment."""

    def test_tabbed_table_header_centered(self):
        """
        Real-world scenario: User creates a centered table header using tabs.
        This mimics what a user would create in Quill after the fix.
        """
        html = '''
        <p class="ql-align-center">Name\tAddress\tPhone</p>
        <p>John Smith\t123 Main St\t555-1234</p>
        <p>Jane Doe\t456 Oak Ave\t555-5678</p>
        '''

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        # Should have 3 paragraphs
        assert len(doc.paragraphs) >= 3

        # First paragraph should be centered
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

        # First paragraph should contain tab characters
        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert text.count('\t') == 2
        assert 'Name' in text and 'Address' in text and 'Phone' in text

        # Other paragraphs should have tabs too
        for i in range(1, 3):
            para_text = ''.join(run.text for run in doc.paragraphs[i].runs)
            assert '\t' in para_text

    def test_right_aligned_with_tab_for_label_value_pairs(self):
        """
        Real-world scenario: Right-aligned label-value pairs with tabs.
        """
        html = '''
        <p style="text-align: right;">Date:\t2024-03-28</p>
        <p style="text-align: right;">Amount:\t$1,000.00</p>
        <p style="text-align: right;">Status:\tPaid</p>
        '''

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        # All paragraphs should be right-aligned with tabs
        for para in doc.paragraphs:
            if para.text.strip():
                assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
                text = ''.join(run.text for run in para.runs)
                assert '\t' in text

    def test_mixed_formatting_preserves_all(self):
        """
        Complex scenario: Mixed alignment, tabs, and other formatting.
        """
        html = '''
        <p class="ql-align-center"><strong>TITLE</strong></p>
        <p>Column1\tColumn2\tColumn3</p>
        <p style="text-align: right;"><em>Right aligned note</em></p>
        <p class="ql-align-justify">This is justified text with\ttabs in the middle.</p>
        '''

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) >= 4

        # Check each paragraph
        # Para 0: centered, bold
        assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert any(run.bold for run in doc.paragraphs[0].runs)

        # Para 1: default alignment with tabs
        text1 = ''.join(run.text for run in doc.paragraphs[1].runs)
        assert text1.count('\t') == 2

        # Para 2: right aligned, italic
        assert doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert any(run.italic for run in doc.paragraphs[2].runs)

        # Para 3: justified with tabs
        assert doc.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        text3 = ''.join(run.text for run in doc.paragraphs[3].runs)
        assert '\t' in text3

    def test_tab_at_start_and_end_of_paragraph(self):
        """
        Edge case: Tabs at the beginning and end of paragraphs.
        """
        html = '''
        <p>\tIndented start</p>
        <p>End with tab\t</p>
        <p>\tBoth sides\t</p>
        '''

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        # Check each paragraph has tabs where expected
        text0 = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert text0.startswith('\t')

        text1 = ''.join(run.text for run in doc.paragraphs[1].runs)
        assert text1.endswith('\t')

        text2 = ''.join(run.text for run in doc.paragraphs[2].runs)
        assert text2.startswith('\t') and text2.endswith('\t')

    def test_alignment_without_tabs_still_works(self):
        """Verify alignment works independently of tabs."""
        html = '''
        <p class="ql-align-left">Left aligned</p>
        <p class="ql-align-center">Center aligned</p>
        <p class="ql-align-right">Right aligned</p>
        <p class="ql-align-justify">Justified text</p>
        '''

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        assert len(doc.paragraphs) >= 4

        # Left is default (None or LEFT)
        assert doc.paragraphs[0].alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)
        assert doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert doc.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert doc.paragraphs[3].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    def test_tabs_without_alignment_still_work(self):
        """Verify tabs work independently of alignment."""
        html = '<p>A\tB\tC\tD\tE</p>'

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(html)

        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert text.count('\t') == 4
        assert text == 'A\tB\tC\tD\tE'
