"""Test template formatting tags: <center>, <right>, <indent>, <tab>."""

import re

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt
from src.services.document_service import DocumentService, HTMLToWordConverter


class TestFormattingTagsInTemplate:
    """Test that formatting tags in templates work correctly."""

    def test_center_tag_basic(self):
        """<center>text</center> should create centered paragraph."""
        template = '<center>TITLE</center>'
        merged = DocumentService._merge_template(template, {}, {})

        # Convert to Word document
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        # Find the paragraph with "TITLE"
        title_para = None
        for para in doc.paragraphs:
            if 'TITLE' in para.text:
                title_para = para
                break

        assert title_para is not None, f"TITLE not found in paragraphs: {[p.text for p in doc.paragraphs]}"
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_center_with_cr_inside(self):
        """<center> with <cr> inside should center all lines (one paragraph, tight line spacing)."""
        template = '<center>LINE1<cr>LINE2<cr>LINE3</center>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        centered = [p for p in doc.paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.text.strip()]
        assert len(centered) >= 1
        block = centered[0].text.replace('\n', '').replace('\r', '')
        assert 'LINE1' in block and 'LINE2' in block and 'LINE3' in block

    def test_right_tag_basic(self):
        """<right>text</right> should create right-aligned paragraph."""
        template = '<right>RIGHT ALIGNED</right>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        right_para = None
        for para in doc.paragraphs:
            if 'RIGHT ALIGNED' in para.text:
                right_para = para
                break

        assert right_para is not None
        assert right_para.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_right_tag_line_shorthand(self):
        """<right>text should right-align only that line through the next <cr>."""
        template = 'Normal line<cr><right>RIGHT LINE<cr>Normal again'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        right_para = None
        normal_after = None
        for para in doc.paragraphs:
            if 'RIGHT LINE' in para.text:
                right_para = para
            if 'Normal again' in para.text:
                normal_after = para

        assert right_para is not None
        assert right_para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert normal_after is not None
        assert normal_after.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT)

    def test_right_with_cr_inside(self):
        """<right> with <cr> inside should right-align all lines in one paragraph."""
        template = '<right>LINE1<cr>LINE2</right>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        right_blocks = [p for p in doc.paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT and p.text.strip()]
        assert len(right_blocks) >= 1
        text = right_blocks[0].text.replace('\n', '').replace('\r', '')
        assert 'LINE1' in text and 'LINE2' in text

    def test_right_preserves_editor_paragraph_breaks_inside_tag(self):
        """Paragraph breaks created inside <right> should remain Word line breaks."""
        template = '<p><right>LINE1</p><p>LINE2</right></p>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        right_blocks = [p for p in doc.paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT and p.text.strip()]
        assert len(right_blocks) == 1
        assert right_blocks[0].text == 'LINE1\nLINE2'

    def test_indent_tag_basic(self):
        """<indent>text</indent> should create indented paragraph."""
        template = '<indent>Indented text</indent>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        # Check that paragraph has left margin/indent
        indent_para = None
        for para in doc.paragraphs:
            if 'Indented text' in para.text:
                indent_para = para
                break

        assert indent_para is not None
        # The paragraph should have left_indent set
        assert indent_para.paragraph_format.left_indent is not None

    def test_indent_with_cr_inside(self):
        """<indent> with <cr> inside should indent all lines in one paragraph."""
        template = '<indent>LINE1<cr>LINE2</indent>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        indented = [p for p in doc.paragraphs if p.paragraph_format.left_indent is not None and p.text.strip()]
        assert len(indented) >= 1
        text = indented[0].text.replace('\n', '').replace('\r', '')
        assert 'LINE1' in text and 'LINE2' in text

    def test_tab_tag_basic(self):
        """<tab> should insert tab character."""
        template = 'Name<tab>Address<tab>Phone'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        # Check that tab characters are present
        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert '\t' in text
        assert text.count('\t') == 2

    def test_center_with_identifier(self):
        """<center> tag should work with identifiers."""
        template = '<center><<title>></center>'
        answer_map = {'title': 'LAST WILL AND TESTAMENT'}
        merged = DocumentService._merge_template(template, answer_map, answer_map)

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        title_para = None
        for para in doc.paragraphs:
            if 'LAST WILL AND TESTAMENT' in para.text:
                title_para = para
                break

        assert title_para is not None
        assert title_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_center_with_identifier_and_cr(self):
        """<center> with identifiers and <cr> inside."""
        template = '<center>LAST WILL<cr>OF<cr><<name>></center>'
        answer_map = {'name': 'John Doe'}
        merged = DocumentService._merge_template(template, answer_map, answer_map)

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        centered = [p for p in doc.paragraphs if p.alignment == WD_ALIGN_PARAGRAPH.CENTER and p.text.strip()]
        assert len(centered) >= 1
        text = centered[0].text.replace('\n', '').replace('\r', '')
        assert 'LAST WILL' in text and 'OF' in text and 'John Doe' in text

    def test_mixed_formatting(self):
        """Multiple formatting tags in one template."""
        template = '''
<center>ARTICLE ##</center>
<cr>
<indent>This is an indented paragraph.</indent>
<cr>
Name<tab>Address<tab>Phone
<cr>
<right>Page ##</right>
'''
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        # Check that we have multiple paragraphs with different formatting
        assert len(doc.paragraphs) > 0

        # Check for centered ARTICLE
        article_found = False
        for para in doc.paragraphs:
            if 'ARTICLE' in para.text:
                assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
                article_found = True
                break
        assert article_found, f"ARTICLE not found in: {[p.text for p in doc.paragraphs]}"

    def test_case_insensitive_tags(self):
        """Tags should work in uppercase, lowercase, or mixed case."""
        templates = [
            '<center>Text</center>',
            '<CENTER>Text</CENTER>',
            '<Center>Text</Center>',
        ]

        for template in templates:
            merged = DocumentService._merge_template(template, {}, {})
            doc = Document()
            parser = HTMLToWordConverter(doc)
            parser.feed(merged)

            centered = False
            for para in doc.paragraphs:
                if 'Text' in para.text:
                    centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
                    break

            assert centered, f"Failed for template: {template}"

    def test_tab_uppercase(self):
        """<TAB> should also work."""
        template = 'A<TAB>B'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        text = ''.join(run.text for run in doc.paragraphs[0].runs)
        assert '\t' in text

    def test_cr_after_center_before_body_no_giant_gap(self):
        """<cr> after </center> must not become an extra paragraph + break (Word gap bug)."""
        template = '<center>TITLE</center><cr>Body paragraph starts here.'
        merged = DocumentService._merge_template(template, {}, {})
        # Should not contain </p><br/> before "Body" (that produced oversized vertical space in Word)
        assert '</p><br/>Body' not in merged.replace(' ', '')
        assert 'Body paragraph' in merged

    def test_no_extra_blank_paragraphs(self):
        """Adjacent <center> blocks merge to one paragraph (tight line spacing in Word)."""
        template = '<center>LINE1</center><center>LINE2</center>'
        merged = DocumentService._merge_template(template, {}, {})

        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)

        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) == 1
        text = non_empty[0].text.replace('\n', '').replace('\r', '')
        assert 'LINE1' in text and 'LINE2' in text
        assert non_empty[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_two_center_blocks_with_cr_merge_tight(self):
        """Title line + name line as separate centers should not double paragraph-gap in Word."""
        template = '<center>DURABLE TITLE</center><cr><center>Person Name</center>'
        merged = DocumentService._merge_template(template, {}, {})
        assert merged.count('ql-align-center') == 1
        assert '<br/>' in merged or 'br' in merged.lower()
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)
        centered = [p for p in doc.paragraphs if p.text.strip() and p.alignment == WD_ALIGN_PARAGRAPH.CENTER]
        assert len(centered) == 1
        t = centered[0].text.replace('\n', '').replace('\r', '')
        assert 'DURABLE TITLE' in t and 'Person Name' in t

    def test_quill_center_alignment_survives_template_flattening(self):
        """Stored Quill alignment must survive even without explicit <center> tags."""
        merged = DocumentService._merge_template(
            '<p class="ql-align-center">AFFIDAVIT OF WITNESSES&lt;cr&gt;</p>'
            '<p class="ql-align-justify">&lt;cr&gt;</p>'
            '<p class="ql-align-center">&lt;center&gt;NOTARY PUBLIC&lt;/center&gt;&lt;cr&gt;</p>',
            {},
            {}
        )

        doc = Document()
        HTMLToWordConverter(doc).feed(merged)

        centered = [
            paragraph for paragraph in doc.paragraphs
            if paragraph.text.strip() and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        ]
        centered_text = "\n".join(paragraph.text for paragraph in centered)
        assert "AFFIDAVIT OF WITNESSES" in centered_text
        assert "NOTARY PUBLIC" in centered_text

    def test_distant_center_blocks_do_not_swallow_intervening_right_blocks(self):
        """Only adjacent aligned paragraphs should be merged."""
        merged = DocumentService._merge_template(
            '<p class="ql-align-center">ACKNOWLEDGMENT OF PRINCIPAL&lt;cr&gt;</p>'
            '<p class="ql-align-justify">Body text&lt;cr&gt;</p>'
            '<p class="ql-align-center">AFFIDAVIT OF WITNESSES&lt;cr&gt;</p>'
            '<p class="ql-align-justify">Witness &lt;right&gt;Witness&lt;/right&gt;</p>'
            '<p class="ql-align-center">&lt;center&gt;NOTARY PUBLIC&lt;/center&gt;&lt;cr&gt;</p>',
            {},
            {}
        )

        doc = Document()
        HTMLToWordConverter(doc).feed(merged)

        notary = next(paragraph for paragraph in doc.paragraphs if "NOTARY PUBLIC" in paragraph.text)
        witness = next(
            paragraph for paragraph in doc.paragraphs
            if paragraph.text.strip() == "Witness" and paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        )
        assert notary.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert witness.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_br_must_not_become_paragraph_breaks_in_word_pipeline(self):
        """Regression: merge_document used to replace every <br/> with </p><p>, duplicating Word
        paragraph spacing for <cr> line breaks. Feeding merged HTML as-is must stay one paragraph."""
        merged = DocumentService._merge_template(
            '<center>DURABLE TITLE</center><cr><center>Person Name</center>', {}, {}
        )
        assert '<br/>' in merged.replace(' ', '')
        broken = re.sub(r'<br\s*/?>', '</p><p>', merged)
        doc_ok = Document()
        HTMLToWordConverter(doc_ok).feed(merged)
        doc_bad = Document()
        HTMLToWordConverter(doc_bad).feed(broken)
        ok_n = len([p for p in doc_ok.paragraphs if p.text.strip()])
        bad_n = len([p for p in doc_bad.paragraphs if p.text.strip()])
        assert ok_n == 1
        assert bad_n >= 2

    def test_word_converter_clears_default_paragraph_spacing(self):
        """Word Normal style adds space before/after paragraphs; merged docs use <cr> for gaps."""
        merged = DocumentService._merge_template(
            '<center>A</center><center>B</center>Body', {}, {}
        )
        doc = Document()
        HTMLToWordConverter(doc).feed(merged)
        for p in doc.paragraphs:
            if not p.text.strip():
                continue
            assert p.paragraph_format.space_before == Pt(0)
            assert p.paragraph_format.space_after == Pt(0)
            assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE

    def test_adjacent_right_blocks_merge_tight(self):
        """Back-to-back <right> blocks should merge like <center> (no double paragraph gap)."""
        template = '<right>R1</right><right>R2</right>'
        merged = DocumentService._merge_template(template, {}, {})
        assert merged.count('ql-align-right') == 1
        doc = Document()
        parser = HTMLToWordConverter(doc)
        parser.feed(merged)
        rights = [p for p in doc.paragraphs if p.text.strip() and p.alignment == WD_ALIGN_PARAGRAPH.RIGHT]
        assert len(rights) == 1
        t = rights[0].text.replace('\n', '').replace('\r', '')
        assert 'R1' in t and 'R2' in t


class TestMergeStackedAlignedParagraphHelpers:
    """Unit tests for _merge_stacked_quill_aligned_paragraphs (used after <center>/<right> expansion)."""

    def test_merge_two_center_paragraphs_html(self):
        html = '<p class="ql-align-center">Line A</p> <p class="ql-align-center">Line B</p>'
        out = DocumentService._merge_stacked_quill_aligned_paragraphs(html)
        assert out.count('ql-align-center') == 1
        assert '<br/>' in out
        assert 'Line A' in out and 'Line B' in out

    def test_merge_three_center_chain(self):
        html = (
            '<p class="ql-align-center">A</p>'
            '<p class="ql-align-center">B</p>'
            '<p class="ql-align-center">C</p>'
        )
        out = DocumentService._merge_stacked_quill_aligned_paragraphs(html)
        assert out.count('ql-align-center') == 1
        assert out.count('<br/>') == 2
        assert 'A' in out and 'B' in out and 'C' in out

    def test_does_not_merge_center_with_left_body(self):
        html = '<p class="ql-align-center">Title</p><p>Body text</p>'
        out = DocumentService._merge_stacked_quill_aligned_paragraphs(html)
        assert out.count('ql-align-center') == 1
        assert 'Title' in out and 'Body text' in out
        assert out.count('<p') == 2
